from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_final_technical_document import (
    choose_diagram_font,
    draw_arrow,
    draw_box,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_cell_width,
    set_repeat_table_header,
    set_run_font,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = OUTPUT_DIR / "competition_paper_assets"
OUTPUT_PATH = OUTPUT_DIR / "面向岗位能力达成的学生成长诊断与精准就业智能体系统_技术论文_竞赛版.docx"

PAPER_TITLE = "面向岗位能力达成的学生成长诊断与精准就业智能体系统"
PAPER_TITLE_EN = (
    "An Intelligent Agent System for Student Growth Diagnosis and Precise Employment "
    "Based on Job Competency Attainment"
)

BODY_CN = "宋体"
HEADING_CN = "黑体"
EN_FONT = "Times New Roman"
INK = "000000"
MUTED = "595959"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
BORDER = "808080"
TABLE_WIDTH = 8760


def exact_20pt(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)


def set_paragraph_keep(paragraph, keep_next: bool = False) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.paragraph_format.widow_control = True


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)

    specs = {
        "Heading 1": (18, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": (14, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (12, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, alignment) in specs.items():
        style = doc.styles[name]
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.page_break_before = False
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = EN_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption.paragraph_format.line_spacing = Pt(18)


def configure_page(section, *, body: bool) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.2)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.clear()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.first_line_indent = Pt(0)
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    header_paragraph.paragraph_format.line_spacing = Pt(14)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.first_line_indent = Pt(0)
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)

    if body:
        set_run_font(
            header_paragraph.add_run("第二十一届中国研究生电子设计竞赛\n"),
            cn=BODY_CN,
            size=9,
            color=INK,
        )
        set_run_font(
            header_paragraph.add_run(PAPER_TITLE),
            cn=BODY_CN,
            size=9,
            color=INK,
        )
        add_field(footer_paragraph, "PAGE", "1")
        for run in footer_paragraph.runs:
            set_run_font(run, cn=BODY_CN, size=9, color=INK)


def restart_page_number(section, start: int = 1) -> None:
    section_pr = section._sectPr
    page_number = section_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_pr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def add_body(doc: Document, text: str, *, first_indent: bool = True, bold_lead: str = ""):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_indent else Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    exact_20pt(paragraph)
    set_paragraph_keep(paragraph)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(paragraph.add_run(bold_lead), cn=HEADING_CN, size=12, bold=True, color=INK)
        set_run_font(paragraph.add_run(text[len(bold_lead):]), cn=BODY_CN, size=12, color=INK)
    else:
        set_run_font(paragraph.add_run(text), cn=BODY_CN, size=12, color=INK)
    return paragraph


def add_chapter(doc: Document, title: str) -> None:
    paragraph = doc.add_heading(title, level=1)
    paragraph.paragraph_format.page_break_before = title != "第1章 绪论"
    paragraph.paragraph_format.space_after = Pt(8)


def add_section(doc: Document, title: str) -> None:
    paragraph = doc.add_heading(title, level=2)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)


def add_subsection(doc: Document, title: str) -> None:
    paragraph = doc.add_heading(title, level=3)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)


def add_table_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    exact_20pt(paragraph)
    set_paragraph_keep(paragraph, keep_next=True)
    set_run_font(paragraph.add_run(text), cn=BODY_CN, size=10.5, color=INK)


def add_source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(16)
    set_run_font(paragraph.add_run(f"资料来源：{text}"), cn=BODY_CN, size=9, color=MUTED)


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    widths: list[int],
    *,
    font_size: float = 9.5,
) -> None:
    add_table_title(doc, title)
    row_values = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, widths, indent_dxa=120)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(16)
        set_run_font(paragraph.add_run(str(header)), cn=HEADING_CN, size=font_size, bold=True, color=INK)

    for row_index, values in enumerate(row_values):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFAFA")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index == 0 and len(headers) > 2
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(18)
            set_run_font(paragraph.add_run(str(value)), cn=BODY_CN, size=font_size, color=INK)


def add_figure(doc: Document, path: Path, caption: str, source: str, width: float = 6.0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    set_paragraph_keep(paragraph, keep_next=True)
    paragraph.add_run().add_picture(str(path), width=Inches(width))

    caption_paragraph = doc.add_paragraph(style="Caption")
    set_paragraph_keep(caption_paragraph, keep_next=True)
    set_run_font(caption_paragraph.add_run(caption), cn=BODY_CN, size=10.5, color=INK)
    add_source_note(doc, source)


def add_formula(doc: Document, formula: str, number: str, note: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    exact_20pt(paragraph)
    set_paragraph_keep(paragraph)
    set_run_font(paragraph.add_run(formula), cn=BODY_CN, size=12, color=INK)
    tab = paragraph.add_run(f"    {number}")
    set_run_font(tab, cn=BODY_CN, size=12, color=INK)
    if note:
        add_body(doc, note, first_indent=False)


def make_collaboration_diagram(path: Path) -> None:
    image = Image.new("RGB", (1650, 920), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    role_font = choose_diagram_font(30, bold=True)
    text_font = choose_diagram_font(23)
    draw.text((70, 35), "学生—学校—企业协同服务架构", font=title_font, fill="#1F4E79")

    draw_box(draw, (80, 180, 450, 410), "学生端\n简历与能力证据\n画像、岗位、路径", fill="#EAF2F8", font=role_font)
    draw_box(draw, (640, 100, 1010, 330), "学校端\n课程与培养方案\n诊断、干预、跟踪", fill="#F2F2F2", font=role_font)
    draw_box(draw, (1200, 180, 1570, 410), "企业端\n真实岗位与技能要求\n岗位反馈、人才筛选", fill="#E2F0D9", font=role_font)
    draw_box(
        draw,
        (410, 540, 1240, 790),
        "精准就业智能体平台\n能力画像工作流｜岗位知识库｜本地快速匹配｜AI 精排｜就业指导｜结构化导出",
        fill="#FFF2CC",
        outline="#806000",
        font=text_font,
    )
    draw_arrow(draw, (265, 420), (540, 540), color="#1F4E79")
    draw_arrow(draw, (825, 340), (825, 530), color="#1F4E79")
    draw_arrow(draw, (1385, 420), (1110, 540), color="#1F4E79")
    draw_arrow(draw, (540, 530), (265, 420), color="#5B9BD5")
    draw_arrow(draw, (1110, 530), (1385, 420), color="#70AD47")
    draw.text(
        (250, 835),
        "平台把学生能力证据、学校培养资源和企业岗位需求转换为可追踪的结构化闭环。",
        font=text_font,
        fill="#595959",
    )
    image.save(path)


def make_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(25, bold=True)
    text_font = choose_diagram_font(21)
    draw.text((65, 35), "单次 LLM、五专家角色与七工具协作工作流", font=title_font, fill="#1F4E79")
    boxes = [
        (55, 160, 340, 330, "画像采集智能体\n文本标准化\n技能扫描"),
        (410, 160, 695, 330, "四维评分智能体\nRubric 审计\n一次 LLM 请求"),
        (765, 160, 1050, 330, "证据抽取智能体\n证据矩阵\n事实边界复核"),
        (1120, 160, 1405, 330, "能力归因智能体\n一致性审计\n发展焦点"),
        (1475, 160, 1745, 330, "质量复核智能体\n最终审计\n结构化输出"),
    ]
    colors = ["#EAF2F8", "#FFF2CC", "#E2F0D9", "#E4DFEC", "#F2F2F2"]
    for index, box in enumerate(boxes):
        draw_box(draw, box[:4], box[4], fill=colors[index], outline="#7F7F7F", font=box_font)
        if index < len(boxes) - 1:
            draw_arrow(draw, (box[2] + 5, 245), (boxes[index + 1][0] - 5, 245), color="#1F4E79")
    draw_box(
        draw,
        (210, 500, 1590, 710),
        "共享工作区\nnormalized_text｜recognized_skills｜ability_scores｜score_evidence｜evidence_cards｜audit_results",
        fill="#F7F7F7",
        outline="#1F4E79",
        font=text_font,
    )
    for x in (197, 552, 907, 1262, 1610):
        draw_arrow(draw, (x, 345), (x, 490), color="#7F7F7F", width=4)
    draw.text(
        (220, 770),
        "执行记录保留 5 个角色、7 次工具调用和 5 次交接；语义生成集中为 1 次 LLM 请求。",
        font=text_font,
        fill="#595959",
    )
    image.save(path)


def make_match_diagram(path: Path) -> None:
    image = Image.new("RGB", (1780, 960), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(24, bold=True)
    text_font = choose_diagram_font(21)
    draw.text((65, 35), "岗位匹配快速路径、AI 增强与按需路径生成", font=title_font, fill="#1F4E79")

    boxes = [
        (50, 150, 300, 300, "读取已保存画像\n复用四维分数"),
        (355, 150, 605, 300, "MySQL 缓存\n优先查询"),
        (660, 150, 910, 300, "向量召回 TOP30\n岗位族限额"),
        (965, 150, 1215, 300, "本地规则评分\n标题归一去重"),
        (1270, 150, 1715, 300, "即时展示 TOP5\n写入 local 缓存"),
    ]
    for index, box in enumerate(boxes):
        draw_box(draw, box[:4], box[4], fill="#EAF2F8", outline="#7F7F7F", font=box_font)
        if index < len(boxes) - 1:
            draw_arrow(draw, (box[2] + 5, 225), (boxes[index + 1][0] - 5, 225), color="#1F4E79")

    lower = [
        (190, 480, 505, 625, "用户按需触发\n本地 TOP10"),
        (575, 480, 890, 625, "一次 AI 双向精排\n120 s 单次超时"),
        (960, 480, 1275, 625, "最多 3 次尝试\n部分结果容错"),
        (1345, 480, 1660, 625, "写入 llm 缓存\n刷新直接复用"),
    ]
    for index, box in enumerate(lower):
        draw_box(draw, box[:4], box[4], fill="#FFF2CC", outline="#BF9000", font=box_font)
        if index < len(lower) - 1:
            draw_arrow(draw, (box[2] + 5, 552), (lower[index + 1][0] - 5, 552), color="#BF9000")
    draw_arrow(draw, (1490, 315), (350, 470), color="#BF9000", width=4)

    draw_box(
        draw,
        (320, 735, 1460, 880),
        "点击单个岗位后再生成成长路径\n技能缺口｜推荐项目｜三阶段学习任务｜结果写回诊断记录",
        fill="#E2F0D9",
        outline="#548235",
        font=text_font,
    )
    image.save(path)


def make_output_diagram(path: Path) -> None:
    image = Image.new("RGB", (1680, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(24, bold=True)
    text_font = choose_diagram_font(21)
    draw.text((65, 35), "结构化输出与数据持久化关系", font=title_font, fill="#1F4E79")
    draw_box(draw, (70, 150, 430, 340), "学生输入\n九类字段\n简历文本/文件", fill="#EAF2F8", font=box_font)
    draw_box(draw, (660, 125, 1020, 365), "诊断与就业智能体\n能力分数、证据、标签\n岗位排序、成长建议", fill="#FFF2CC", font=box_font)
    draw_box(draw, (1250, 150, 1610, 340), "前端呈现\n雷达图、TOP5\n趋势图、指导建议", fill="#E2F0D9", font=box_font)
    draw_arrow(draw, (440, 245), (650, 245), color="#1F4E79")
    draw_arrow(draw, (1030, 245), (1240, 245), color="#1F4E79")
    draw_box(
        draw,
        (150, 535, 690, 760),
        "MySQL 持久化\n诊断记录｜岗位库｜就业指导记录\nlocal/llm 匹配缓存",
        fill="#F2F2F2",
        font=text_font,
    )
    draw_box(
        draw,
        (990, 535, 1530, 760),
        "结构化导出\n能力画像 JSON\nTOP5 岗位与路径 PDF",
        fill="#F2F2F2",
        font=text_font,
    )
    draw_arrow(draw, (840, 380), (530, 525), color="#7F7F7F")
    draw_arrow(draw, (840, 380), (1150, 525), color="#7F7F7F")
    image.save(path)


def make_performance_chart(path: Path) -> None:
    image = Image.new("RGB", (1650, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    label_font = choose_diagram_font(24, bold=True)
    text_font = choose_diagram_font(21)
    draw.text((70, 35), "岗位匹配关键路径性能实测", font=title_font, fill="#1F4E79")
    labels = ["MySQL缓存读取", "本地匹配平均", "本地匹配P95", "本地匹配最大值"]
    values = [0.368, 90.73, 91.41, 156.75]
    colors = ["#70AD47", "#5B9BD5", "#4472C4", "#ED7D31"]
    max_value = 170.0
    chart_left = 350
    chart_right = 1530
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = 180 + index * 150
        draw.text((75, y + 18), label, font=label_font, fill="#000000")
        width = int((chart_right - chart_left) * value / max_value)
        draw.rounded_rectangle(
            (chart_left, y, chart_left + max(width, 5), y + 72),
            radius=16,
            fill=color,
        )
        draw.text((chart_left + width + 18, y + 18), f"{value:.3f} ms" if value < 1 else f"{value:.2f} ms", font=text_font, fill="#000000")
    draw.text(
        (350, 800),
        "测试对象：679 条岗位记录；本地匹配 30 次，缓存读取 100 次。",
        font=text_font,
        fill="#595959",
    )
    image.save(path)


def create_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "collaboration": ASSET_DIR / "collaboration_architecture.png",
        "workflow": ASSET_DIR / "single_llm_workflow.png",
        "match": ASSET_DIR / "job_match_pipeline.png",
        "output": ASSET_DIR / "structured_output.png",
        "performance": ASSET_DIR / "performance_benchmark.png",
    }
    make_collaboration_diagram(paths["collaboration"])
    make_workflow_diagram(paths["workflow"])
    make_match_diagram(paths["match"])
    make_output_diagram(paths["output"])
    make_performance_chart(paths["performance"])
    return paths


def add_cover(doc: Document) -> None:
    configure_page(doc.sections[0], body=False)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(35)

    competition = doc.add_paragraph()
    competition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    competition.paragraph_format.first_line_indent = Pt(0)
    competition.paragraph_format.space_after = Pt(18)
    set_run_font(
        competition.add_run("第二十一届中国研究生电子设计竞赛"),
        cn=HEADING_CN,
        size=18,
        bold=True,
        color=INK,
    )

    paper_type = doc.add_paragraph()
    paper_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paper_type.paragraph_format.first_line_indent = Pt(0)
    paper_type.paragraph_format.space_after = Pt(36)
    set_run_font(paper_type.add_run("技 术 论 文"), cn=HEADING_CN, size=26, bold=True, color=INK)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    set_run_font(title.add_run(PAPER_TITLE), cn=HEADING_CN, size=18, bold=True, color=INK)

    title_en = doc.add_paragraph()
    title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_en.paragraph_format.first_line_indent = Pt(0)
    title_en.paragraph_format.space_after = Pt(56)
    set_run_font(title_en.add_run(PAPER_TITLE_EN), cn=BODY_CN, en=EN_FONT, size=12, bold=True, color=INK)

    metadata = [
        ("参赛单位", "桂林电子科技大学"),
        ("参赛队名", "智聘未来"),
        ("指导教师", "陈辉  副教授"),
        ("参赛队员", "黎盛、冯靖原、侯程予"),
        ("完成时间", "2026 年 6 月"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(10)
        exact_20pt(paragraph)
        set_run_font(paragraph.add_run(f"{label}："), cn=HEADING_CN, size=15, bold=True, color=INK)
        set_run_font(paragraph.add_run(value), cn=HEADING_CN, size=15, bold=True, color=INK)


def add_abstracts(doc: Document) -> None:
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(front, body=False)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    exact_20pt(heading)
    set_run_font(heading.add_run("摘  要"), cn=HEADING_CN, size=16, bold=True, color=INK)

    abstract = (
        "针对高校学生成长诊断中评价依据分散、岗位推荐同质化、远程大模型调用耗时长以及结果难以复用等问题，"
        "本文设计并实现面向岗位能力达成的学生成长诊断与精准就业智能体系统。系统以 FastAPI 为应用框架，"
        "使用 SQLAlchemy 与 MySQL 管理用户、诊断、岗位知识、就业指导和岗位匹配缓存数据，以 LangGraph 编排"
        "画像采集、四维评分、证据抽取、能力归因和质量复核五类专家角色。为兼顾智能体协作展示与响应效率，"
        "系统将多个专家的语义生成合并为一次大模型请求，同时保留共享工作区、七次工具调用和五次交接记录，"
        "形成可审计的专业基础、技术实践、工具技能和职业发展四维画像。在岗位匹配阶段，系统复用数据库中的画像分数，"
        "通过本地稀疏向量召回、技能覆盖率、项目证据、目标方向与学历适配度完成初排，并使用岗位标题归一化、岗位族配额"
        "和分数阈值控制推荐多样性；用户可按需触发一次 TOP10 大模型双向精排，综合学生适岗分与岗位适生分。"
        "系统进一步将匹配结果按诊断记录、岗位库版本、算法版本和结果类型写入 MySQL 持久化缓存，使应用重启后仍可复用。"
        "更新后的系统还增加精准就业指导、十年职业趋势、能力画像 JSON 导出和 TOP5 岗位路径 PDF 导出。"
        "在 679 条真实岗位数据上，本地匹配 30 次平均耗时 90.73 ms、P95 为 91.41 ms，缓存读取 100 次平均耗时"
        "0.368 ms；岗位匹配与单次 LLM 工作流相关的 11 项自动化测试全部通过。结果表明，该混合架构在保持可解释性、"
        "结构化输出和多智能体协作特征的同时，能够将岗位页面首屏稳定控制在 3 s 指标以内，并改善结果多样性与系统可复现性。"
    )
    add_body(doc, abstract)
    keyword = doc.add_paragraph()
    keyword.paragraph_format.first_line_indent = Pt(0)
    exact_20pt(keyword)
    set_run_font(keyword.add_run("关键词："), cn=HEADING_CN, size=12, bold=True, color=INK)
    set_run_font(
        keyword.add_run("学生能力画像；多智能体协作；岗位匹配；大语言模型；持久化缓存"),
        cn=BODY_CN,
        size=12,
        color=INK,
    )

    doc.add_page_break()
    heading_en = doc.add_paragraph()
    heading_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_en.paragraph_format.first_line_indent = Pt(0)
    heading_en.paragraph_format.space_after = Pt(12)
    exact_20pt(heading_en)
    set_run_font(heading_en.add_run("ABSTRACT"), cn=HEADING_CN, en=EN_FONT, size=16, bold=True, color=INK)

    abstract_en = (
        "This paper presents an intelligent agent system for student growth diagnosis and precise employment "
        "based on job competency attainment. The system addresses fragmented evaluation evidence, homogeneous "
        "job recommendations, long remote-model latency, and the lack of reusable results. FastAPI is used as "
        "the application framework, while SQLAlchemy and MySQL manage users, diagnosis records, job knowledge, "
        "employment guidance, and persistent matching caches. LangGraph orchestrates five expert roles for profile "
        "collection, four-dimensional scoring, evidence extraction, attribution analysis, and quality review. "
        "To preserve visible multi-agent collaboration without repeated semantic calls, the workflow consolidates "
        "generation into one large-language-model request and retains a shared workspace, seven tool calls, and five "
        "handoff records. The job-matching stage reuses stored ability scores and performs local sparse-vector recall, "
        "interpretable rule scoring, title normalization, job-family quotas, and score-threshold diversification. "
        "Users may optionally invoke a single TOP10 bidirectional reranking request that combines student-to-job fit "
        "and job-to-student growth suitability. Results are stored in a version-aware MySQL cache and remain available "
        "after application restart. The updated system also provides precise employment guidance, ten-year career "
        "projection, JSON export for ability profiles, and PDF export for TOP5 job paths. On 679 real job records, "
        "thirty local matching runs averaged 90.73 ms with a P95 of 91.41 ms, while one hundred persistent-cache reads "
        "averaged 0.368 ms. All eleven automated tests related to the core matching path and the single-LLM workflow "
        "passed. The results demonstrate that the hybrid architecture improves responsiveness, explainability, "
        "recommendation diversity, and reproducibility while retaining explicit agent collaboration."
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    exact_20pt(paragraph)
    set_run_font(paragraph.add_run(abstract_en), cn=BODY_CN, en=EN_FONT, size=11, color=INK)

    keywords_en = doc.add_paragraph()
    keywords_en.paragraph_format.first_line_indent = Pt(0)
    exact_20pt(keywords_en)
    set_run_font(keywords_en.add_run("Key words: "), en=EN_FONT, size=11, bold=True, color=INK)
    set_run_font(
        keywords_en.add_run(
            "student ability profile; multi-agent collaboration; job matching; large language model; persistent cache"
        ),
        en=EN_FONT,
        size=11,
        color=INK,
    )


def add_toc(doc: Document) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Pt(0)
    heading.paragraph_format.space_after = Pt(12)
    exact_20pt(heading)
    set_run_font(heading.add_run("目  录"), cn=HEADING_CN, size=16, bold=True, color=INK)

    entries = [
        ("第1章 绪论", 1, 0),
        ("1.1 研究背景与问题", 1, 1),
        ("1.1.1 高校就业指导的数字化需求", 1, 2),
        ("1.1.2 需要解决的关键问题", 1, 2),
        ("1.2 研究目标与内容", 1, 1),
        ("1.2.1 建设目标", 1, 2),
        ("1.2.2 主要研究内容", 1, 2),
        ("第2章 核心创新与协同价值", 2, 0),
        ("2.1 单次 LLM 的多角色协作", 2, 1),
        ("2.1.1 创新机制", 2, 2),
        ("2.1.2 相对优势", 2, 2),
        ("2.2 快速、可靠且多样的岗位匹配", 3, 1),
        ("2.2.1 本地首屏与按需智能增强", 3, 2),
        ("2.2.2 去重与岗位族多样化", 3, 2),
        ("2.3 版本感知的持久化缓存", 3, 1),
        ("2.4 就业指导与结构化交付", 3, 1),
        ("第3章 系统总体架构", 4, 0),
        ("3.1 学生—学校—企业协同架构", 4, 1),
        ("3.2 系统分层与模块划分", 4, 1),
        ("3.2.1 分层架构", 4, 2),
        ("3.2.2 功能模块", 5, 2),
        ("3.3 数据流与部署关系", 5, 1),
        ("第4章 关键技术实现", 6, 0),
        ("4.1 学生数据处理与能力画像", 6, 1),
        ("4.1.1 简历字段提取与文本归一", 6, 2),
        ("4.1.2 四维评分与证据审计", 6, 2),
        ("4.2 岗位召回、评分与 AI 精排", 6, 1),
        ("4.2.1 本地候选召回", 6, 2),
        ("4.2.2 本地可解释评分", 7, 2),
        ("4.2.3 双向 AI 精排", 7, 2),
        ("4.3 精准就业指导与长期发展预测", 8, 1),
        ("第5章 结构化输出与数据设计", 9, 0),
        ("5.1 结构化字段与导出", 9, 1),
        ("5.1.1 能力画像结构", 9, 2),
        ("5.1.2 岗位匹配与成长路径结构", 9, 2),
        ("5.1.3 就业指导结构", 9, 2),
        ("5.2 MySQL 数据模型", 10, 1),
        ("5.3 MySQL 持久化缓存", 10, 1),
        ("第6章 性能测试与实测案例", 11, 0),
        ("6.1 测试环境与方法", 11, 1),
        ("6.1.1 环境与数据", 11, 2),
        ("6.1.2 自动化测试方法", 11, 2),
        ("6.2 性能、稳定性与准确性", 11, 1),
        ("6.2.1 性能结果", 11, 2),
        ("6.2.2 稳定性与测试通过情况", 12, 2),
        ("6.2.3 准确性与可解释性评价", 13, 2),
        ("6.3 学生与企业案例", 13, 1),
        ("6.3.1 学生案例 S-01", 13, 2),
        ("6.3.2 企业岗位案例 E-01", 14, 2),
        ("第7章 商业化方案", 15, 0),
        ("7.1 用户、规模与盈利模式", 15, 1),
        ("7.1.1 目标用户", 15, 2),
        ("7.1.2 应用规模", 15, 2),
        ("7.1.3 盈利模式", 15, 2),
        ("7.2 推广路径与风险控制", 16, 1),
        ("第8章 总结与展望", 17, 0),
        ("8.1 工作总结", 17, 1),
        ("8.2 局限性与后续工作", 17, 1),
        ("参考文献", 18, 0),
    ]
    for title, page, level in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Cm(0.55 * level)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        exact_20pt(paragraph)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.0),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        set_run_font(
            paragraph.add_run(f"{title}\t{page}"),
            cn=HEADING_CN if level == 0 else BODY_CN,
            size=14 if level == 0 else (12 if level == 1 else 10.5),
            bold=level == 0,
            color=INK,
        )


def add_chapter_1(doc: Document) -> None:
    add_chapter(doc, "第1章 绪论")
    add_section(doc, "1.1 研究背景与问题")
    add_subsection(doc, "1.1.1 高校就业指导的数字化需求")
    add_body(
        doc,
        "高校就业指导正在从一次性的简历修改和岗位推荐，转向贯穿学习、实践、求职和入职准备的连续服务。"
        "学生信息通常分散在简历、课程成绩、项目经历、竞赛证书和自我介绍中，教师需要在有限时间内完成能力判断、"
        "岗位解释和成长建议，企业则希望看到与岗位要求直接关联的证据。若缺少统一的数据结构和评价方法，学生容易得到"
        "“技能很多但证据不足”“推荐岗位名称相近但差异不清”等低价值结果。"
    )
    add_body(
        doc,
        "大语言模型具备理解非结构化文本和生成解释性建议的能力，但直接将完整流程串联为多次远程调用，会受到网络、"
        "模型限流、输出格式和费用的影响。项目早期流程曾出现画像生成慢、岗位页面长时间刷新和精排超时等问题。"
        "因此，本系统将可确定计算留在本地，将复杂语义判断限制在少量、可审计、可降级的环节，并以数据库记录支撑结果复用。"
    )

    add_subsection(doc, "1.1.2 需要解决的关键问题")
    add_body(
        doc,
        "第一，如何既体现多智能体工作流和工具协作，又避免多个 LLM 串行调用造成的高延迟；第二，如何把学生当前能力、"
        "岗位要求和未来成长价值统一到可解释的岗位匹配模型；第三，如何解决岗位库中同义标题和相同岗位族重复占据 TOP5；"
        "第四，如何让本地结果、AI 精排和成长路径在应用重启后仍可复用；第五，如何以稳定字段输出能力画像、岗位差距和就业建议，"
        "使前端展示、导出和后续分析不依赖自由文本。"
    )

    add_section(doc, "1.2 研究目标与内容")
    add_subsection(doc, "1.2.1 建设目标")
    add_body(
        doc,
        "系统以“能力证据可追溯、岗位匹配可解释、页面响应可接受、智能增强可按需”为总体目标。能力画像阶段输出专业基础、"
        "技术实践、工具技能和职业发展四维分数及证据；岗位匹配阶段要求首屏不依赖 LLM 并稳定满足 3 s 响应指标；"
        "AI 精排只处理本地 TOP10，成长路径只针对用户点击的单个岗位生成；所有关键结果通过 MySQL 持久化。"
    )
    add_subsection(doc, "1.2.2 主要研究内容")
    add_body(
        doc,
        "本文围绕五项内容展开：一是基于 LangGraph 的单次 LLM 五角色画像工作流；二是基于稀疏向量召回、规则评分、"
        "标题归一化和岗位族多样化的快速匹配算法；三是学生适岗分与岗位适生分融合的双向 AI 精排；四是版本感知的"
        "MySQL 持久化缓存和结构化导出；五是面向学生、学校和企业的协同服务及商业化落地方案。"
    )


def add_chapter_2(doc: Document, diagrams: dict[str, Path]) -> None:
    add_chapter(doc, "第2章 核心创新与协同价值")
    add_section(doc, "2.1 单次 LLM 的多角色协作")
    add_subsection(doc, "2.1.1 创新机制")
    add_body(
        doc,
        "系统不是简单删除多智能体，而是把语义生成集中到一次请求，再通过工具节点完成复核。画像采集智能体调用"
        "ProfileTextNormalizer 和 SkillKeywordScanner；四维评分智能体先调用 RubricScoreCalculator，随后由"
        "LLMAbilityScorer 一次生成四维分数、证据、优势、短板和质量提示；证据抽取、能力归因和质量复核智能体继续调用"
        "EvidenceMatrixTool、ConsistencyAuditTool 和 FinalAuditTool。最终执行记录保留五个专家角色、七次工具调用、"
        "五次交接和共享工作区，能够在页面中完整展示工作流。"
    )
    add_figure(
        doc,
        diagrams["workflow"],
        "图2-1 单次 LLM、五专家角色与七工具协作工作流",
        "作者根据 app/agent/diagnosis_agent.py 绘制。",
        6.15,
    )
    add_subsection(doc, "2.1.2 相对优势")
    add_body(
        doc,
        "与每个专家独立调用模型相比，该方式减少了网络往返和重复上下文，避免四个评分专家对同一份简历重复解释；"
        "与单一提示词直接输出最终结论相比，该方式又保留了评分规程、证据矩阵和一致性审计，能够指出信息不足和结论边界。"
        "自动化测试 test_one_llm_call_keeps_full_workflow_and_tools 验证了“一次 LLM 调用仍保留完整工作流和工具链”。"
    )

    add_section(doc, "2.2 快速、可靠且多样的岗位匹配")
    add_subsection(doc, "2.2.1 本地首屏与按需智能增强")
    add_body(
        doc,
        "岗位匹配页面优先读取 AI 精排缓存，其次读取本地缓存；两者均未命中时，执行不创建 LLM 客户端的本地匹配并立即保存。"
        "用户主动点击“AI 精排”后，系统才把本地 TOP10 一次性发送给模型。用户点击某个岗位后，才为该岗位生成差距与成长路径。"
        "这一分离机制把不可控的远程耗时从页面首屏移出，同时保留 AI 对复杂语义和成长价值的判断能力。"
    )
    add_subsection(doc, "2.2.2 去重与岗位族多样化")
    add_body(
        doc,
        "系统把大小写、空格、括号说明、职级和常见后缀统一为稳定标题键，并将岗位归入应用开发、数据、云运维、测试、"
        "前端、算法、安全等岗位族。TOPN 选择先在高相关候选中每族选取一个，再放宽到每族两个，最后按原分数补齐。"
        "多样性阈值要求岗位至少达到 60 分且与最高分差不超过 18 分；本地路径放宽为最低 40 分、最高分差不超过 25 分。"
        "因此系统不会为了形式上的多样性强行插入低相关岗位。"
    )

    add_section(doc, "2.3 版本感知的持久化缓存")
    add_body(
        doc,
        "传统内存缓存随进程退出而清空。本系统新增 job_match_cache_records 表，分别保存 local 和 llm 两类结果。"
        "缓存键由诊断记录 ID、岗位库版本、算法版本和结果类型共同计算；岗位数量或岗位技能字段变化时，岗位库 SHA-256"
        "版本随之变化；评分公式或提示词升级时，只需修改算法版本即可自动避开旧结果。该机制既提高重复访问速度，也避免"
        "错误复用过期缓存。"
    )

    add_section(doc, "2.4 就业指导与结构化交付")
    add_body(
        doc,
        "更新后的系统新增精准就业指导分支。它从简历中本地提取姓名、专业、学历、技能、项目、证书和竞赛信息，复用四维评分"
        "与岗位库选择 TOP1 发展锚点，生成 0—1 年、1—3 年、3—5 年和 5—10 年建议，并输出职位层级、岗位胜任力、"
        "技术深度、项目影响力和职业成熟度趋势。能力画像可导出 JSON，岗位 TOP5 与路径可导出 PDF，使结果能够进入"
        "学生档案、教师复核和答辩材料。"
    )


def add_chapter_3(doc: Document, diagrams: dict[str, Path]) -> None:
    add_chapter(doc, "第3章 系统总体架构")
    add_section(doc, "3.1 学生—学校—企业协同架构")
    add_body(
        doc,
        "学生端提供简历、项目、技能和求职意向，获得能力画像、岗位推荐和发展路径；学校端提供课程、培养方案和就业指导，"
        "根据结构化证据识别共性短板并进行干预；企业端提供真实岗位名称、技能要求、学历、城市和薪资等信息，形成岗位能力"
        "知识库。平台把三方数据转换为统一字段和可解释结果，使学生成长诊断不再脱离企业需求，也使企业岗位要求能够反向"
        "支持学校课程和实践项目调整。"
    )
    add_figure(
        doc,
        diagrams["collaboration"],
        "图3-1 学生—学校—企业协同服务架构",
        "作者根据系统业务角色与数据流绘制。",
        6.1,
    )

    add_section(doc, "3.2 系统分层与模块划分")
    add_subsection(doc, "3.2.1 分层架构")
    add_body(
        doc,
        "系统采用服务端渲染的分层单体架构。表示层由 Jinja2、HTML、CSS 和 JavaScript 构成，负责表单、雷达图、岗位列表和"
        "趋势图；应用层由 FastAPI 路由、Session 登录校验、文件上传和模板渲染构成；智能与算法层包含画像工作流、"
        "四维规则评分、岗位召回与精排、就业指导、简历优化、课程映射和成长路径；数据层由 MySQL 和本地岗位 Excel 数据组成。"
    )
    add_table(
        doc,
        "表3-1 系统技术选型",
        ["层次", "技术", "选型依据"],
        [
            ("表示层", "Jinja2、HTML、CSS、JavaScript", "依赖少，便于竞赛部署和服务端会话控制"),
            ("应用层", "FastAPI、Starlette Session", "路由清晰，可同时提供页面与 JSON 接口"),
            ("工作流", "LangGraph", "共享状态适合表达多节点协作与工具调用"),
            ("模型接入", "LangChain ChatOpenAI", "兼容 DeepSeek 等 OpenAI 协议接口"),
            ("数据层", "SQLAlchemy 2.x、MySQL", "支持声明式模型、索引与持久化 JSON"),
            ("检索排序", "稀疏向量、余弦相似度、规则评分", "本地可解释且无需额外向量数据库"),
        ],
        [1200, 2500, 5060],
    )
    add_source_note(doc, "作者根据当前项目依赖与代码实现整理。")

    add_subsection(doc, "3.2.2 功能模块")
    add_table(
        doc,
        "表3-2 核心功能模块",
        ["模块", "主要输入", "主要输出"],
        [
            ("学生信息与简历解析", "表单、PDF/DOCX/TXT 简历", "九类学生字段与标准化文本"),
            ("多智能体能力画像", "学生字段、简历全文", "四维分数、证据卡、标签、风险和工作流日志"),
            ("岗位匹配", "画像、岗位知识记录", "本地 TOP10、AI 精排 TOP10、TOP5 多样化结果"),
            ("成长路径", "学生信息、单个岗位", "缺口、项目建议和三阶段行动计划"),
            ("精准就业指导", "简历文本、岗位库", "TOP1 锚点、就业建议和十年趋势"),
            ("结构化导出", "画像与岗位结果", "能力画像 JSON、TOP5 岗位路径 PDF"),
        ],
        [2300, 2900, 3560],
    )
    add_source_note(doc, "作者根据 app/main.py 的 28 个 HTTP 路由及服务模块整理。")

    add_section(doc, "3.3 数据流与部署关系")
    add_body(
        doc,
        "浏览器通过 FastAPI 页面路由提交简历和指令，应用创建数据库会话并读取当前用户最近一次诊断。对可确定任务，服务直接"
        "调用本地函数；对语义任务，服务通过配置文件中的模型地址、密钥和模型名访问兼容接口。MySQL 保存用户、诊断、"
        "岗位、就业指导和缓存记录。当前主应用定义 28 个 HTTP 路由，代码覆盖登录、首页智能体、就业指导、简历优化、"
        "简历匹配、模拟面试、学生信息、画像、历史、岗位、健康检查和成长趋势等页面。"
    )


def add_chapter_4(doc: Document, diagrams: dict[str, Path]) -> None:
    add_chapter(doc, "第4章 关键技术实现")
    add_section(doc, "4.1 学生数据处理与能力画像")
    add_subsection(doc, "4.1.1 简历字段提取与文本归一")
    add_body(
        doc,
        "系统将学生信息统一为 name、major、grade、target_job、skills、projects、competitions、certificates 和 self_intro"
        "九类字段。首页智能体根据指令识别画像、简历优化或就业指导分支；就业指导使用本地正则和技术词表提取字段，画像分支"
        "使用结构化简历抽取服务。针对 PDF 文本层把英文拆成“J a v a”“M y S Q L”的情况，关键词检测先压缩连续英文字符间的"
        "空格再匹配，从而减少技术词漏识别。"
    )
    add_subsection(doc, "4.1.2 四维评分与证据审计")
    add_body(
        doc,
        "四维评分分别关注专业课程与基础知识、项目和实习的实践深度、语言框架与工程工具、求职目标与表达准备。"
        "RubricScoreCalculator 提供本地评分参考，LLM 必须返回 ability_scores 和 score_evidence；EvidenceMatrixTool"
        "把证据转为四维矩阵，ConsistencyAuditTool 检查高分但证据不足等冲突，FinalAuditTool 复核工具链完整性。"
        "诊断结果保存到 diagnosis_records.agent_result_json，页面只展示画像相关章节，不在画像页混入岗位排序。"
    )
    add_table(
        doc,
        "表4-1 四维能力定义",
        ["维度", "评价内容", "主要证据"],
        [
            ("专业基础", "课程、专业知识、计算机基础、相关证书", "数据结构、算法、数据库、网络、学历"),
            ("技术实践", "项目复杂度、职责、部署、优化和结果", "项目、实习、竞赛、量化成果"),
            ("工具技能", "语言、框架、数据库、中间件和工程工具", "Java、Spring、MySQL、Redis、Docker、Git"),
            ("职业发展", "目标清晰度、材料完整性、表达和面试准备", "目标岗位、自我介绍、证书、沟通协作"),
        ],
        [1500, 3500, 3760],
    )
    add_source_note(doc, "作者根据 app/services/ability_match_service.py 和画像提示词整理。")

    add_section(doc, "4.2 岗位召回、评分与 AI 精排")
    add_subsection(doc, "4.2.1 本地候选召回")
    add_body(
        doc,
        "岗位文本由岗位名、技能、项目、课程、证书、学历、薪资和描述组成。学生查询文本由目标岗位、专业、技能、项目和证书组成。"
        "中文连续片段生成 2—4 字 n-gram，英文按单词切分，词频采用对数缩放；系统计算学生向量与岗位向量的余弦相似度，"
        "先召回较大候选池，再按岗位族限额形成最多 30 条方向分散的候选。"
    )
    add_formula(doc, "w(t)=1+ln(tf(t))", "（4-1）")
    add_formula(doc, "sim(q,d)=(q·d)/(||q||₂×||d||₂)", "（4-2）")

    add_subsection(doc, "4.2.2 本地可解释评分")
    add_body(
        doc,
        "本地评分综合技能覆盖、项目证据、四维画像、岗位方向、学历和证书。技能覆盖占主要权重，避免仅凭岗位名称相似就获得高分；"
        "项目证据反映候选人是否在真实场景使用过相关技术；岗位方向和学历提供适配修正。每个结果同时输出已匹配技能、缺失技能、"
        "相关项目和推荐理由，便于学生理解得分来源。"
    )
    add_formula(
        doc,
        "M_local=55C_skill+15C_project+15A_profile+10F_role+3F_edu+2C_cert",
        "（4-3）",
    )
    add_formula(
        doc,
        "A_profile=0.30P+0.25R+0.30T+0.15C",
        "（4-4）",
        "其中 P、R、T、C 分别表示专业基础、技术实践、工具技能和职业发展能力的归一化分数。",
    )

    add_subsection(doc, "4.2.3 双向 AI 精排")
    add_body(
        doc,
        "AI 精排只接收本地 TOP10，并要求一次返回每个岗位的学生适岗分、岗位适生分、最终分和不超过 35 个汉字的理由。"
        "学生适岗分衡量当前技能、项目、学历和职业方向是否满足岗位；岗位适生分衡量缺口可补齐性、难度适配、学习路径可映射性"
        "和成长价值。最终分按 60% 与 40% 加权。单次请求默认超时 120 s，最多尝试 3 次并采用指数退避；若部分条目解析失败，"
        "系统保留成功结果并用本地候选补齐，而页面始终保留本地排序。"
    )
    add_formula(doc, "M_AI=0.60S_student→job+0.40S_job→student", "（4-5）")
    add_figure(
        doc,
        diagrams["match"],
        "图4-1 岗位匹配快速路径、AI 增强与按需路径生成",
        "作者根据 app/services/llm_ability_match_service.py 与 app/main.py 绘制。",
        6.15,
    )

    add_section(doc, "4.3 精准就业指导与长期发展预测")
    add_body(
        doc,
        "就业指导分支不依赖远程 LLM。系统先对简历进行本地字段抽取和四维评分，再使用岗位规则匹配选择 TOP1 发展锚点。"
        "随后生成入职、第 1 年、第 3 年、第 5 年、第 7 年和第 10 年六个节点，并分别预测职位层级、岗位胜任力、技术深度、"
        "项目影响力和职业成熟度。建议内容按就业主线定位、简历证据强化、短板补齐、面试准备以及四个职业阶段组织，"
        "每条建议均附依据和行动项，避免只输出宏观口号。"
    )


def add_chapter_5(doc: Document, diagrams: dict[str, Path]) -> None:
    add_chapter(doc, "第5章 结构化输出与数据设计")
    add_section(doc, "5.1 结构化字段与导出")
    add_subsection(doc, "5.1.1 能力画像结构")
    add_body(
        doc,
        "能力画像导出负载包含学生基础字段、四维分数、雷达图数据、画像总结、标签、优势、短板、维度洞察、证据卡、发展焦点、"
        "风险提示和质量复核；同时附带 TOP5 岗位、匹配来源、是否命中缓存、算法版本，以及智能体角色、工作流步骤、工具调用、"
        "协作日志和复核发现。该 JSON 可作为教师复核、后续分析和第三方系统集成的统一接口。"
    )
    add_subsection(doc, "5.1.2 岗位匹配与成长路径结构")
    add_body(
        doc,
        "每个岗位结果包含岗位名、公司、城市、薪资、学历要求、匹配分、已匹配技能、缺失技能、相关项目、推荐理由、"
        "学生适岗分、岗位适生分和岗位族。成长路径包含差距清单、推荐项目和学习阶段；每个阶段包括阶段名、持续时间、目标、"
        "行动任务和阶段成果。TOP5 导出 PDF 将上述字段整理为学生信息表、岗位汇总表和逐岗位路径明细。"
    )
    add_subsection(doc, "5.1.3 就业指导结构")
    add_body(
        doc,
        "就业指导记录保存 student、ability_scores、summary、top1_job、evidence_basis、precision_guidance、"
        "development_suggestions、trend 和 trend_chart。trend_chart 中每条序列都使用统一的六个时间节点，前端可直接绘制"
        "多折线趋势，无需再次解析自然语言。"
    )
    add_figure(
        doc,
        diagrams["output"],
        "图5-1 结构化输出与数据持久化关系",
        "作者根据新增导出、就业指导和缓存代码绘制。",
        6.05,
    )

    add_section(doc, "5.2 MySQL 数据模型")
    add_table(
        doc,
        "表5-1 主要业务表与当前数据快照",
        ["数据表", "用途", "当前记录"],
        [
            ("users", "用户注册、登录与创建时间", "1"),
            ("diagnosis_records", "学生九类字段、四维分数与智能体 JSON", "18"),
            ("job_knowledge_records", "岗位、公司、城市、学历、薪资、技能和项目", "679"),
            ("job_match_cache_records", "本地与 AI 匹配结果的持久化缓存", "5"),
            ("employment_guidance_records", "就业指导输入与结构化结果", "0"),
        ],
        [2600, 4500, 1660],
    )
    add_source_note(doc, "2026 年 6 月 12 日当前 MySQL 实测快照；数据由项目配置的 MySQL 实例实时查询获得。")
    add_body(
        doc,
        "当前岗位库包含 278 个不同岗位名称、643 个不同公司名称和 175 个不同城市字段值。岗位知识记录以 JSON 文本保存技能、"
        "项目、推荐课程和推荐证书，便于在不频繁改变表结构的情况下扩展岗位知识。诊断记录按 user_id 和创建时间查询最近结果，"
        "历史页面保留多次诊断，从而支持成长趋势分析。"
    )

    add_section(doc, "5.3 MySQL 持久化缓存")
    add_subsection(doc, "5.3.1 缓存键与索引")
    add_body(
        doc,
        "job_match_cache_records.cache_key 建立唯一索引，diagnosis_id 建立普通索引。job_version 对岗位 ID、岗位名、技能、"
        "项目、课程、证书、学历、经验、路径和薪资等字段计算 SHA-256；algorithm_version 当前为"
        "match_cache_v7_reliable_ai_refine；result_type 区分 local 与 llm。"
    )
    add_formula(
        doc,
        "cache_key=SHA-256(diagnosis_id:job_version:algorithm_version:result_type)",
        "（5-1）",
    )
    add_subsection(doc, "5.3.2 重启后仍有效的原因")
    add_body(
        doc,
        "应用重启只清空 Python 进程内存，不会删除 MySQL 表中的 result_json。只要重新启动后仍连接同一数据库，且诊断记录、"
        "岗位库版本和算法版本没有变化，系统即可构造相同 cache_key 并读取结果。因此“重启后仍有效”本质上是磁盘数据库持久化"
        "与版本化键共同作用；若数据库卷被删除、岗位数据变化或算法版本升级，系统会自动重新计算。"
    )


def add_chapter_6(doc: Document, diagrams: dict[str, Path]) -> None:
    add_chapter(doc, "第6章 性能测试与实测案例")
    add_section(doc, "6.1 测试环境与方法")
    add_subsection(doc, "6.1.1 环境与数据")
    add_body(
        doc,
        "测试基于 2026 年 6 月 12 日更新后的主分支代码，后端运行于 Python 3.11 环境，数据库为 MySQL，岗位表包含 679 条记录。"
        "本地性能测试直接复用最新诊断记录和数据库中的四维分数，避免将网络波动混入本地算法指标。性能脚本连续执行 30 次本地"
        "匹配和 100 次持久化缓存读取，统计平均值、P95 和最大值。计时使用 Python time.perf_counter。"
    )
    add_subsection(doc, "6.1.2 自动化测试方法")
    add_body(
        doc,
        "项目使用 unittest 和 mock 验证单次 LLM 工作流、岗位标题归一化、候选池岗位族限额、TOP5 多样化、无关岗位不强制插入、"
        "本地快速路径不创建 LLM 客户端、TOP10 一次性精排、成长路径本地降级和 PDF 拆分技术词识别。完整测试集共 16 项。"
    )

    add_section(doc, "6.2 性能、稳定性与准确性")
    add_subsection(doc, "6.2.1 性能结果")
    add_table(
        doc,
        "表6-1 岗位匹配关键路径性能",
        ["测试项", "样本数", "平均耗时", "P95", "最大值", "结果"],
        [
            ("MySQL 持久化缓存读取", "100", "0.368 ms", "0.554 ms", "未单列", "通过"),
            ("679 岗位本地匹配", "30", "90.73 ms", "91.41 ms", "156.75 ms", "通过"),
            ("岗位页面首屏指标", "上述路径", "< 3 s", "< 3 s", "< 3 s", "通过"),
            ("历史 TOP10 AI 精排", "1", "约 73.92 s", "不适用", "不适用", "按需任务"),
        ],
        [2200, 900, 1450, 1300, 1400, 1510],
        font_size=9,
    )
    add_source_note(doc, "本地与缓存数据为本次实测；AI 精排为优化阶段保留的真实端到端观测值。")
    add_figure(
        doc,
        diagrams["performance"],
        "图6-1 岗位匹配关键路径性能实测",
        "作者根据 30 次本地匹配与 100 次缓存读取结果绘制。",
        6.0,
    )
    add_body(
        doc,
        "结果表明，页面首屏路径远低于 3 s 指标。AI 精排仍受远程模型影响，但它已被移到用户主动触发的后台交互中，"
        "且完成后写入 MySQL，后续刷新不再重复调用。论文将 AI 精排视为可选增强任务，而不是首屏服务等级的一部分。"
    )

    add_subsection(doc, "6.2.2 稳定性与测试通过情况")
    add_body(
        doc,
        "本地匹配 30 次和缓存读取 100 次均成功返回，未出现空结果和异常。完整自动化测试 16 项中 12 项通过；"
        "其中岗位匹配、岗位多样性、快速路径和单次 LLM 工作流相关的 11 项全部通过。4 项失败集中在简历优化服务"
        "response_format 参数契约变化，以及简历抽取测试未隔离外部网络而误触真实模型，不能据此宣称完整测试集全部通过。"
    )
    add_table(
        doc,
        "表6-2 自动化测试分类结果",
        ["测试类别", "通过/总数", "结论"],
        [
            ("单次 LLM 多智能体工作流", "1/1", "一次调用仍保留完整工作流与工具"),
            ("岗位多样性与标题归一", "4/4", "同义岗位去重和岗位族控制有效"),
            ("岗位快速路径与降级", "6/6", "默认路径不创建 LLM 客户端，失败可本地回退"),
            ("简历优化与字段抽取", "1/5", "4 项需同步测试契约并隔离外部依赖"),
            ("完整测试集", "12/16", "当前总体通过率 75%"),
        ],
        [2700, 1500, 4560],
    )
    add_source_note(doc, "2026 年 6 月 12 日执行 python -m unittest discover -s tests -v 的结果。")

    add_subsection(doc, "6.2.3 准确性与可解释性评价")
    add_body(
        doc,
        "项目尚未建立带人工相关性标签的大规模岗位推荐金标准，因此本文不虚构“模型准确率”。当前采用三类工程代理指标："
        "标题归一化与岗位族单元测试通过率 100%；本地 TOP5 中岗位标题归一键无重复；所有岗位结果均包含已匹配技能、缺失技能、"
        "得分组件和推荐理由。下一阶段应邀请就业教师和企业招聘人员对推荐相关性、理由可信度和路径可执行性进行双盲标注，"
        "再计算 Precision@5、NDCG@5 和专家一致性。"
    )

    add_section(doc, "6.3 学生与企业案例")
    add_subsection(doc, "6.3.1 学生案例 S-01")
    add_body(
        doc,
        "案例 S-01 取自当前数据库最新诊断记录并做匿名化处理。学生专业为地理空间信息工程，学历为本科，目标岗位为 Java 后端工程师。"
        "画像分数为专业基础 63、技术实践 66、工具技能 40、职业发展 68。证据包括物流运输管理微服务项目、Java、Spring Boot、"
        "Spring Cloud、MySQL、Redis、Docker、Git、Maven 和 Linux；主要风险是非计算机专业理论基础和企业级项目验证不足。"
    )
    add_table(
        doc,
        "表6-3 案例 S-01 的本地 TOP5 推荐",
        ["序号", "岗位", "匹配分", "岗位族", "主要待补技能"],
        [
            ("1", "软件开发工程师", "66", "应用开发", "Oracle、Elasticsearch"),
            ("2", "数据库开发工程师（暑期实习）", "58", "数据库", "PostgreSQL、ClickHouse"),
            ("3", "云计算工程师", "58", "云运维", "云主机、Shell"),
            ("4", "JAVA 后端开发工程师", "54", "应用开发", "Kubernetes、Kafka、Jenkins"),
            ("5", "后端开发（大数据研究工程师）", "53", "数据", "Hadoop、Hive、Spark"),
        ],
        [700, 2900, 1000, 1300, 2860],
    )
    add_source_note(doc, "当前 MySQL 诊断记录 ID=19 与 679 条岗位库的本地匹配结果。")
    add_body(
        doc,
        "案例结果不再被多个 Java 同义岗位完全占据，而是覆盖应用开发、数据库、云运维和数据四个方向。第一岗位与学生已掌握的"
        "Java、SpringMVC、MySQL、Redis 和消息队列形成直接证据；第二至第五岗位体现可迁移方向，帮助学生比较稳定就业、"
        "数据库深化、云平台运维和大数据后端等不同路径。"
    )

    add_subsection(doc, "6.3.2 企业岗位案例 E-01")
    add_body(
        doc,
        "企业案例 E-01 对应岗位知识记录 ID=622，岗位为北京爱回收科技有限公司的软件开发工程师，工作城市北京，薪资"
        "25—35K·14 薪，学历要求本科。岗位技能包括 Java、SpringMVC、MySQL、Oracle、Redis、消息队列和 Elasticsearch。"
        "系统识别 S-01 已匹配 Java、SpringMVC、MySQL、Redis 和消息队列，缺口为 Oracle 与 Elasticsearch，项目证据命中交易场景，"
        "本地得分 66。企业可据此快速看到候选人“已具备什么、缺什么、证据来自哪里”，学生也能把补齐任务转化为可验证项目。"
    )


def add_chapter_7(doc: Document) -> None:
    add_chapter(doc, "第7章 商业化方案")
    add_section(doc, "7.1 用户、规模与盈利模式")
    add_subsection(doc, "7.1.1 目标用户")
    add_body(
        doc,
        "第一类用户为高校就业指导中心、学院和专业负责人，用于学生能力普查、重点群体识别和培养方案反馈；第二类用户为学生，"
        "用于个人画像、岗位匹配、简历优化、面试训练和成长跟踪；第三类用户为企业招聘与校企合作部门，用于岗位能力标准化、"
        "候选人初筛和定向人才培养；第四类用户为职业教育机构和培训机构，用于课程—能力—岗位闭环验证。"
    )
    add_subsection(doc, "7.1.2 应用规模")
    add_body(
        doc,
        "产品可从单学院 500—3000 名学生的私有化试点开始，随后扩展到校级就业平台。岗位库采用版本化导入，画像和匹配结果"
        "以用户和诊断记录隔离。对于跨校部署，可按学校设置独立数据库或租户字段；对于企业合作，可只开放匿名画像和岗位差距，"
        "避免直接暴露学生原始简历。"
    )
    add_subsection(doc, "7.1.3 盈利模式")
    add_table(
        doc,
        "表7-1 商业化产品与收费方式",
        ["产品形态", "主要客户", "收费方式", "核心价值"],
        [
            ("校级 SaaS", "高校就业中心", "按年度与学生规模订阅", "统一画像、岗位推荐和数据看板"),
            ("私有化部署", "高校、职业院校、大型企业", "软件许可、部署与运维服务费", "数据留校、可对接现有系统"),
            ("企业人才专场", "招聘企业", "按岗位包、活动或成功推荐收费", "岗位标准化与候选人能力证据"),
            ("学生增值服务", "个人学生", "基础免费，深度报告与训练订阅", "简历、面试、路径和趋势跟踪"),
            ("知识库服务", "学校与企业", "岗位库、课程能力库更新服务费", "保持岗位技能与培养内容同步"),
        ],
        [1800, 1800, 2500, 2660],
    )
    add_source_note(doc, "作者结合系统模块与高校就业服务场景设计。")

    add_section(doc, "7.2 推广路径与风险控制")
    add_subsection(doc, "7.2.1 推广路径")
    add_body(
        doc,
        "第一阶段以一个学院完成数据接入、教师复核和学生使用闭环，建立画像与推荐的人工标注集；第二阶段接入课程、实习和就业去向，"
        "形成专业级能力差距分析；第三阶段与企业共建岗位能力模板和项目题库，开展定向培养与双选；第四阶段将匿名统计能力开放给"
        "区域高校和产业联盟，形成岗位变化与人才供给趋势服务。"
    )
    add_subsection(doc, "7.2.2 主要风险及应对")
    add_table(
        doc,
        "表7-2 商业化风险与控制措施",
        ["风险", "表现", "控制措施"],
        [
            ("隐私与合规", "简历含个人信息，跨机构流转风险高", "最小化采集、字段脱敏、访问控制、审计日志和数据留校"),
            ("模型幻觉", "模型可能生成无证据结论", "强制结构化字段、证据矩阵、本地审计和人工复核"),
            ("岗位数据时效", "岗位要求和薪资变化快", "记录来源、更新时间和岗位库版本，定期增量导入"),
            ("推荐偏差", "单一专业或岗位族过度占优", "标题归一、岗位族多样化、人工标注与公平性评估"),
            ("外部服务波动", "LLM 超时或接口不可用", "本地首屏、重试、部分容错、持久化缓存和降级结果"),
            ("工程安全", "当前比赛账号密码为明文", "正式部署前采用密码哈希、强 Session 密钥、HTTPS 和数据库迁移"),
        ],
        [1500, 3000, 4260],
    )
    add_source_note(doc, "作者根据当前实现边界与生产部署要求整理。")


def add_chapter_8(doc: Document) -> None:
    add_chapter(doc, "第8章 总结与展望")
    add_section(doc, "8.1 工作总结")
    add_body(
        doc,
        "本文完成了一套从简历与学生信息采集、四维能力画像、岗位快速匹配、AI 双向精排、单岗位成长路径到精准就业指导和"
        "结构化导出的完整系统。项目在架构上用一次 LLM 请求保留五角色协作，在算法上用本地召回与规则评分保障首屏，"
        "在推荐层用标题归一和岗位族控制降低同质化，在数据层用 MySQL 版本化缓存保证重启后复用。当前 679 条岗位数据上的"
        "本地匹配平均耗时 90.73 ms，缓存读取平均耗时 0.368 ms，核心流程测试全部通过。"
    )
    add_section(doc, "8.2 局限性与后续工作")
    add_body(
        doc,
        "系统仍有三类工作需要继续推进。首先，建立由就业教师和企业专家共同标注的岗位相关性数据集，使用 Precision@5、"
        "NDCG@5 和专家一致性替代工程代理指标。其次，修复简历优化 response_format 测试契约，给简历抽取测试注入假模型，"
        "恢复模拟面试服务导入，并增加就业指导和导出功能的自动化测试。最后，生产部署应完成密码哈希、租户隔离、HTTPS、"
        "数据库迁移、模型调用审计、岗位来源追踪和个人信息脱敏。"
    )
    add_body(
        doc,
        "后续还可引入经过审核的课程—能力—岗位图谱、校内项目与企业实习数据，使系统从“根据简历判断能力”进一步发展为"
        "“根据学习过程持续积累能力证据”；同时将 AI 精排改为任务队列和异步通知，使用户在网络不稳定时仍能获得最终精排结果。"
    )


def add_references(doc: Document) -> None:
    add_chapter(doc, "参考文献")
    references = [
        "[1] FastAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/.",
        "[2] SQLAlchemy Authors. SQLAlchemy 2.0 Documentation[EB/OL]. https://docs.sqlalchemy.org/.",
        "[3] LangChain AI. LangGraph Documentation[EB/OL]. https://langchain-ai.github.io/langgraph/.",
        "[4] Oracle Corporation. MySQL Reference Manual[EB/OL]. https://dev.mysql.com/doc/.",
        "[5] Manning C D, Raghavan P, Schütze H. Introduction to Information Retrieval[M]. Cambridge University Press, 2008.",
        "[6] Carbonell J, Goldstein J. The Use of MMR, Diversity-Based Reranking for Reordering Documents and Producing Summaries[C]//SIGIR. 1998.",
        "[7] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//EMNLP-IJCNLP. 2019.",
        "[8] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS. 2020.",
        "[9] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]//NeurIPS. 2017.",
        "[10] 中华人民共和国主席令. 中华人民共和国个人信息保护法[Z]. 2021.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.hanging_indent = Pt(24)
        exact_20pt(paragraph)
        set_run_font(paragraph.add_run(reference), cn=BODY_CN, en=EN_FONT, size=10.5, color=INK)


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()
    doc = Document()
    configure_styles(doc)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_cover(doc)
    add_abstracts(doc)
    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section, body=True)
    restart_page_number(body_section, 1)

    add_chapter_1(doc)
    add_chapter_2(doc, diagrams)
    add_chapter_3(doc, diagrams)
    add_chapter_4(doc, diagrams)
    add_chapter_5(doc, diagrams)
    add_chapter_6(doc, diagrams)
    add_chapter_7(doc)
    add_chapter_8(doc)
    add_references(doc)

    doc.core_properties.title = PAPER_TITLE
    doc.core_properties.subject = "第二十一届中国研究生电子设计竞赛技术论文"
    doc.core_properties.author = "智聘未来参赛队"
    doc.core_properties.keywords = "学生能力画像, 多智能体, 岗位匹配, 大语言模型, MySQL持久化缓存"
    doc.core_properties.comments = "基于 2026-06-12 更新后的项目代码、数据库快照和测试结果生成"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
