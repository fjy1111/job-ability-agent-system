from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = OUTPUT_DIR / "technical_document_assets"
OUTPUT_PATH = OUTPUT_DIR / "岗位能力达成学生成长诊断与精准就业智能体系统_技术文档_最终版.docx"

BODY_CN = "宋体"
HEADING_CN = "黑体"
EN_FONT = "Times New Roman"
MONO_FONT = "Consolas"
INK = "111827"
NAVY = "17365D"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
PALE_BLUE = "F5F8FD"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"
GREEN = "166534"
AMBER = "92400E"
RED = "B42318"


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 0) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)
                set_cell_margins(row.cells[index])
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    *,
    cn: str = BODY_CN,
    en: str = EN_FONT,
    size: float = 12,
    bold: bool | None = None,
    color: str = INK,
    italic: bool | None = None,
) -> None:
    run.font.name = en
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), en)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), en)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_paragraph_shading(paragraph, fill: str, border_color: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "4")
        element.set(qn("w:color"), border_color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.5

    heading_specs = {
        "Heading 1": (16, 16, 8),
        "Heading 2": (14, 12, 6),
        "Heading 3": (12, 8, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.2

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
        style.font.size = Pt(12)
        style.paragraph_format.left_indent = Cm(0.74)
        style.paragraph_format.first_line_indent = Cm(-0.37)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5

    caption = doc.styles["Caption"]
    caption.font.name = EN_FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    caption.font.size = Pt(10.5)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def configure_section(section, *, cover: bool = False) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.first_line_indent = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    if not cover:
        set_run_font(
            header_para.add_run("岗位能力达成学生成长诊断与精准就业智能体系统技术文档"),
            cn=BODY_CN,
            size=9,
            color=GRAY,
        )
        p_pr = header_para._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), BORDER)
        borders.append(bottom)
        p_pr.append(borders)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.first_line_indent = Pt(0)
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)
    if not cover:
        set_run_font(footer_para.add_run("第 "), cn=BODY_CN, size=9, color=GRAY)
        add_field(footer_para, "PAGE", "1")
        set_run_font(footer_para.add_run(" 页"), cn=BODY_CN, size=9, color=GRAY)


def restart_page_number(section, start: int = 1) -> None:
    section_pr = section._sectPr
    pg_num = section_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def add_body_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold_prefix: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    first_indent: bool = True,
    color: str = INK,
    after: float = 5,
) -> object:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_indent else Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.5

    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), bold=True, color=color)
        set_run_font(paragraph.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run_font(paragraph.add_run(text), color=color)
    return paragraph


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.74 + level * 0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.37)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.5
    set_run_font(paragraph.add_run(text))


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.5
    set_run_font(paragraph.add_run(text))


def add_formula(doc: Document, formula: str, explanation: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run(formula), cn=BODY_CN, en=EN_FONT, size=12, bold=True, color=NAVY)
    if explanation:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.first_line_indent = Pt(0)
        note.paragraph_format.space_after = Pt(7)
        set_run_font(note.add_run(explanation), size=10.5, color=GRAY)


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_BLUE, color: str = NAVY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.35
    set_paragraph_shading(paragraph, fill)
    set_run_font(paragraph.add_run(f"{label}："), bold=True, color=color)
    set_run_font(paragraph.add_run(text), color=color)


def add_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    widths: list[int],
    *,
    font_size: float = 10.5,
    header_fill: str = LIGHT_BLUE,
) -> object:
    row_values = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        set_run_font(paragraph.add_run(str(header)), cn=HEADING_CN, size=font_size, bold=True, color=NAVY)

    for row_index, values in enumerate(row_values):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.18
            set_run_font(paragraph.add_run(str(value)), size=font_size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    set_run_font(paragraph.add_run(text), size=10.5, color=GRAY)


def add_image(doc: Document, path: Path, caption: str, width_inches: float = 6.0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def choose_diagram_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_path = (
        "/System/Library/Fonts/STHeiti Medium.ttc"
        if bold
        else "/System/Library/Fonts/Hiragino Sans GB.ttc"
    )
    return ImageFont.truetype(font_path, size=size)


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    lines = text.split("\n")
    boxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    heights = [item[3] - item[1] for item in boxes]
    total_height = sum(heights) + max(0, len(lines) - 1) * 9
    y = box[1] + (box[3] - box[1] - total_height) / 2
    for line, bounds, height in zip(lines, boxes, heights):
        width = bounds[2] - bounds[0]
        x = box[0] + (box[2] - box[0] - width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += height + 9


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    *,
    fill: str = "#F5F8FD",
    outline: str = "#94A3B8",
    font=None,
    radius: int = 18,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, font or choose_diagram_font(30), "#17365D")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#2563EB", width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 18, y2 - 11), (x2 - direction * 18, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 11, y2 - direction * 18), (x2 + 11, y2 - direction * 18)]
    draw.polygon(points, fill=color)


def make_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 940), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(29, bold=True)
    small_font = choose_diagram_font(24)
    draw.text((70, 35), "系统总体分层架构", font=title_font, fill="#17365D")

    layers = [
        ((80, 125, 1520, 245), "表示层", "#EAF2FF", [
            ((150, 155, 405, 215), "Jinja2 页面", small_font),
            ((480, 155, 735, 215), "能力雷达图", small_font),
            ((810, 155, 1065, 215), "岗位匹配页", small_font),
            ((1140, 155, 1450, 215), "简历/面试/成长页面", small_font),
        ]),
        ((80, 290, 1520, 430), "应用服务层", "#F0FDF4", [
            ((140, 325, 390, 395), "FastAPI 路由", small_font),
            ((450, 325, 710, 395), "认证与会话", small_font),
            ((770, 325, 1030, 395), "业务编排", small_font),
            ((1090, 325, 1460, 395), "文件解析与模板渲染", small_font),
        ]),
        ((80, 475, 1520, 635), "智能与算法层", "#FFF7ED", [
            ((120, 515, 365, 600), "LangGraph\n多智能体工作流", small_font),
            ((415, 515, 660, 600), "本地四维评分\n与证据规则", small_font),
            ((710, 515, 955, 600), "稀疏向量召回\n余弦相似度", small_font),
            ((1005, 515, 1250, 600), "岗位去重\n与多样化", small_font),
            ((1300, 515, 1490, 600), "LLM 精排\n与路径生成", small_font),
        ]),
        ((80, 680, 1520, 850), "数据与基础设施层", "#F8FAFC", [
            ((135, 725, 405, 815), "MySQL\n业务数据与持久化缓存", small_font),
            ((485, 725, 755, 815), "岗位知识库\n679 条岗位快照", small_font),
            ((835, 725, 1105, 815), "OpenAI 兼容\nLLM 服务", small_font),
            ((1185, 725, 1460, 815), "静态资源\nExcel / PDF / DOCX", small_font),
        ]),
    ]

    for outer, label, fill, inner_boxes in layers:
        draw.rounded_rectangle(outer, radius=22, fill=fill, outline="#94A3B8", width=3)
        draw.text((outer[0] + 18, outer[1] + 10), label, font=box_font, fill="#17365D")
        for box, text, font in inner_boxes:
            draw_box(draw, box, text, fill="#FFFFFF", outline="#CBD5E1", font=font, radius=14)

    for y1, y2 in ((245, 290), (430, 475), (635, 680)):
        draw_arrow(draw, (800, y1 + 4), (800, y2 - 4), color="#2563EB", width=5)

    image.save(path)


def make_agent_workflow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1700, 780), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(25, bold=True)
    small_font = choose_diagram_font(21)
    draw.text((65, 32), "能力画像单次 LLM 协作与五智能体工作流", font=title_font, fill="#17365D")

    boxes = [
        (70, 165, 345, 330, "01 画像采集智能体\n文本标准化 + 技能扫描"),
        (405, 165, 680, 330, "02 四维评分智能体\nRubric 工具 + 1 次 LLM"),
        (740, 165, 1015, 330, "03 证据抽取智能体\n证据矩阵复核"),
        (1075, 165, 1350, 330, "04 能力归因智能体\n一致性审计"),
        (1410, 165, 1660, 330, "05 质量复核智能体\n最终审计"),
    ]
    colors = ["#EAF2FF", "#FFF7ED", "#F0FDF4", "#F5F3FF", "#F8FAFC"]
    for index, (x1, y1, x2, y2, text) in enumerate(boxes):
        draw_box(draw, (x1, y1, x2, y2), text, fill=colors[index], outline="#64748B", font=box_font)
        if index < len(boxes) - 1:
            draw_arrow(draw, (x2 + 8, 247), (boxes[index + 1][0] - 8, 247))

    workspace = (190, 470, 1510, 655)
    draw_box(
        draw,
        workspace,
        "共享工作区 Shared Workspace\nnormalized_text、ability_scores、score_evidence、evidence_cards、audit_results",
        fill="#F8FAFC",
        outline="#2563EB",
        font=small_font,
    )
    for x in (207, 542, 877, 1212, 1535):
        draw_arrow(draw, (x, 340), (x, 462), color="#64748B", width=4)

    draw.text(
        (225, 690),
        "工作流保留 5 个专家角色、7 次工具调用和 5 次交接记录，但语义生成仅使用 1 次 LLM 请求。",
        font=small_font,
        fill="#475467",
    )
    image.save(path)


def make_match_pipeline_diagram(path: Path) -> None:
    image = Image.new("RGB", (1700, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(25, bold=True)
    small_font = choose_diagram_font(21)
    draw.text((60, 35), "岗位匹配快速路径与按需 AI 增强", font=title_font, fill="#17365D")

    top_boxes = [
        (55, 150, 300, 290, "读取最新诊断\n与四维画像"),
        (355, 150, 600, 290, "MySQL 缓存\n优先查询"),
        (655, 150, 900, 290, "本地稀疏向量\n召回候选"),
        (955, 150, 1200, 290, "规则评分\n与岗位族去重"),
        (1255, 150, 1645, 290, "立即展示 TOP5\n保存 local 缓存"),
    ]
    for index, box in enumerate(top_boxes):
        draw_box(draw, box[:4], box[4], fill="#EAF2FF" if index < 2 else "#F0FDF4", outline="#64748B", font=box_font)
        if index < len(top_boxes) - 1:
            draw_arrow(draw, (box[2] + 8, 220), (top_boxes[index + 1][0] - 8, 220))

    draw.text((70, 345), "用户按需触发", font=box_font, fill="#92400E")
    lower_boxes = [
        (155, 410, 445, 555, "本地 TOP10\n一次性发送"),
        (520, 410, 810, 555, "AI 双向精排\n120 秒上限"),
        (885, 410, 1175, 555, "最多 3 次重试\n部分结果容错"),
        (1250, 410, 1560, 555, "保存 llm 缓存\n刷新直接命中"),
    ]
    for index, box in enumerate(lower_boxes):
        draw_box(draw, box[:4], box[4], fill="#FFF7ED", outline="#D97706", font=box_font)
        if index < len(lower_boxes) - 1:
            draw_arrow(draw, (box[2] + 8, 482), (lower_boxes[index + 1][0] - 8, 482), color="#D97706")
    draw_arrow(draw, (1450, 300), (310, 400), color="#D97706", width=4)

    callout = (170, 665, 1530, 825)
    draw_box(
        draw,
        callout,
        "实测结果\n真实 TOP10 AI 精排约 73.92 s；完成后写入 MySQL；刷新读取缓存约 178 ms。",
        fill="#F8FAFC",
        outline="#2563EB",
        font=small_font,
    )
    image.save(path)


def make_cache_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(image)
    title_font = choose_diagram_font(42, bold=True)
    box_font = choose_diagram_font(25, bold=True)
    small_font = choose_diagram_font(21)
    draw.text((60, 35), "MySQL 持久化缓存键与失效机制", font=title_font, fill="#17365D")

    inputs = [
        (80, 160, 350, 275, "诊断记录 ID"),
        (450, 160, 720, 275, "岗位库版本\nSHA-256"),
        (820, 160, 1090, 275, "算法版本"),
        (1190, 160, 1460, 275, "结果类型\nlocal / llm"),
    ]
    for box in inputs:
        draw_box(draw, box[:4], box[4], fill="#EAF2FF", outline="#64748B", font=box_font)
        draw_arrow(draw, ((box[0] + box[2]) // 2, 285), (770, 390), color="#2563EB")

    draw_box(
        draw,
        (535, 405, 1005, 540),
        "cache_key = SHA-256\n(diagnosis_id : job_version : algorithm_version : result_type)",
        fill="#FFF7ED",
        outline="#D97706",
        font=small_font,
    )
    draw_arrow(draw, (770, 550), (770, 645), color="#166534")
    draw_box(
        draw,
        (405, 655, 1135, 765),
        "job_match_cache_records\n唯一索引查询 + result_json 持久化；应用重启后仍可读取",
        fill="#F0FDF4",
        outline="#16A34A",
        font=small_font,
    )
    image.save(path)


def create_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "architecture": ASSET_DIR / "architecture.png",
        "workflow": ASSET_DIR / "agent_workflow.png",
        "match_pipeline": ASSET_DIR / "match_pipeline.png",
        "cache": ASSET_DIR / "persistent_cache.png",
    }
    make_architecture_diagram(paths["architecture"])
    make_agent_workflow_diagram(paths["workflow"])
    make_match_pipeline_diagram(paths["match_pipeline"])
    make_cache_diagram(paths["cache"])
    return paths


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, cover=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(54)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.first_line_indent = Pt(0)
    kicker.paragraph_format.space_after = Pt(22)
    set_run_font(kicker.add_run("项目技术论文式说明书"), cn=HEADING_CN, size=18, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    set_run_font(
        title.add_run("岗位能力达成学生成长诊断与\n精准就业智能体系统"),
        cn=HEADING_CN,
        size=28,
        bold=True,
        color=NAVY,
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Pt(0)
    subtitle.paragraph_format.space_after = Pt(72)
    set_run_font(subtitle.add_run("设计、实现与性能优化技术文档（最终版）"), cn=HEADING_CN, size=16, color=GRAY)

    metadata = [
        ("系统名称", "job-ability-agent-system"),
        ("文档性质", "基于最终工作区代码的技术论文式说明"),
        ("基准日期", "2026 年 6 月 12 日"),
        ("技术栈", "FastAPI · LangGraph · SQLAlchemy · MySQL · LLM"),
        ("数据快照", "679 条岗位数据 · 18 条诊断记录 · 5 条匹配缓存"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 6880])
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for index, text in enumerate((label, value)):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            set_run_font(
                paragraph.add_run(text),
                cn=HEADING_CN if index == 0 else BODY_CN,
                size=11,
                bold=index == 0,
                color=NAVY if index == 0 else INK,
            )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.first_line_indent = Pt(0)
    note.paragraph_format.space_before = Pt(48)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("编制日期：2026 年 6 月 12 日"), size=12, color=GRAY)

    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Pt(0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(18)
    set_run_font(heading.add_run("摘  要"), cn=HEADING_CN, size=18, bold=True, color=NAVY)

    chinese_abstract = (
        "面向高校学生就业能力评价中存在的信息分散、评价依据不透明、岗位推荐同质化以及大模型调用时延较高等问题，"
        "本文设计并实现了一套岗位能力达成学生成长诊断与精准就业智能体系统。系统采用 FastAPI 构建 Web 应用，"
        "使用 SQLAlchemy 与 MySQL 管理用户、诊断、岗位知识和缓存数据，并以 LangGraph 编排学生能力画像工作流。"
        "在能力诊断阶段，系统保留画像采集、四维评分、证据抽取、能力归因和质量复核五类专家角色，通过共享工作区、"
        "工具调用日志和智能体交接记录体现协作过程，同时将原有多次语义生成压缩为一次大模型请求，以降低画像生成耗时。"
        "在岗位匹配阶段，系统先利用本地稀疏词袋向量和余弦相似度召回候选岗位，再依据技能覆盖、项目证据、画像分数、"
        "目标方向和学历适配度进行可解释评分，并通过岗位标题归一化和岗位族配额缓解推荐结果同质化。用户可按需触发一次"
        "大模型双向精排，综合学生适岗分与岗位适生分生成最终排序。为避免重复计算，系统新增 MySQL 持久化缓存，"
        "以诊断记录、岗位库版本、算法版本和结果类型构造缓存键，保证项目重启后仍可复用本地排序和 AI 精排结果。"
        "当前数据库快照包含 679 条岗位知识记录。真实接口测试中，TOP10 AI 精排约耗时 73.92 s，完成后页面刷新并命中"
        "持久化缓存约耗时 178 ms。核心岗位匹配相关的 10 项自动化测试全部通过。结果表明，该混合架构能够在保持"
        "可解释性和智能体协作展示能力的同时，显著改善岗位匹配入口的响应速度、稳定性和结果多样性。"
    )
    add_body_paragraph(doc, chinese_abstract)
    keywords = doc.add_paragraph()
    keywords.paragraph_format.first_line_indent = Pt(0)
    keywords.paragraph_format.space_before = Pt(8)
    keywords.paragraph_format.space_after = Pt(16)
    set_run_font(keywords.add_run("关键词："), cn=HEADING_CN, bold=True)
    set_run_font(keywords.add_run("学生能力画像；多智能体；岗位匹配；大语言模型；持久化缓存；可解释推荐"))

    heading_en = doc.add_paragraph()
    heading_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_en.paragraph_format.first_line_indent = Pt(0)
    heading_en.paragraph_format.space_before = Pt(14)
    heading_en.paragraph_format.space_after = Pt(14)
    set_run_font(heading_en.add_run("ABSTRACT"), cn=HEADING_CN, en=EN_FONT, size=16, bold=True, color=NAVY)

    english_abstract = (
        "This document presents the design and implementation of an intelligent system for student ability diagnosis, "
        "job matching, and personalized growth planning. The system is built with FastAPI, SQLAlchemy, MySQL, and "
        "LangGraph. A five-role diagnosis workflow preserves explicit agent collaboration, tool calls, evidence review, "
        "and handoff logs, while semantic generation is consolidated into a single large-language-model request. "
        "For job recommendation, a local sparse-vector retrieval stage and an interpretable scoring model first produce "
        "fast candidates. Title normalization and job-family quotas are then applied to reduce redundant recommendations. "
        "An optional LLM reranking stage evaluates both student-to-job fit and job-to-student growth suitability. "
        "A MySQL-backed persistent cache stores local and LLM results with version-aware SHA-256 keys, allowing cached "
        "results to survive application restarts. The current database snapshot contains 679 job records. In an observed "
        "end-to-end run, reranking ten jobs required approximately 73.92 seconds, while a subsequent cached page refresh "
        "required approximately 178 milliseconds. The implemented hybrid architecture improves responsiveness, "
        "explainability, reliability, and recommendation diversity."
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.5
    set_run_font(paragraph.add_run(english_abstract), cn=BODY_CN, en=EN_FONT, size=11.5)
    keywords_en = doc.add_paragraph()
    keywords_en.paragraph_format.first_line_indent = Pt(0)
    keywords_en.paragraph_format.space_after = Pt(0)
    set_run_font(keywords_en.add_run("Key words: "), en=EN_FONT, size=11.5, bold=True)
    set_run_font(
        keywords_en.add_run(
            "student ability profile; multi-agent workflow; job matching; large language model; persistent cache; explainable recommendation"
        ),
        en=EN_FONT,
        size=11.5,
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Pt(0)
    heading.paragraph_format.space_after = Pt(16)
    set_run_font(heading.add_run("目  录"), cn=HEADING_CN, size=18, bold=True, color=NAVY)

    toc_entries = [
        ("摘要", 1, 0),
        ("ABSTRACT", 1, 0),
        ("1 绪论", 4, 0),
        ("2 需求分析", 5, 0),
        ("3 系统总体设计", 6, 0),
        ("4 多智能体能力画像设计", 9, 0),
        ("5 岗位匹配算法设计", 11, 0),
        ("6 数据库与持久化缓存设计", 13, 0),
        ("7 其他业务模块设计", 15, 0),
        ("8 性能优化与可靠性设计", 16, 0),
        ("9 系统实现与接口", 18, 0),
        ("10 测试与验证", 19, 0),
        ("11 部署、运维与安全", 21, 0),
        ("12 局限性与后续工作", 22, 0),
        ("结论", 22, 0),
        ("参考文献", 23, 0),
        ("附录 A 关键源代码索引", 23, 0),
        ("附录 B 最终版本验收检查表", 24, 0),
        ("附录 C 文档基准说明", 24, 0),
    ]
    for title, page, level in toc_entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Cm(0.55 * level)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(14.6),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        set_run_font(paragraph.add_run(f"{title}\t{page}"), size=11.5, bold=level == 0)
    doc.add_page_break()


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("1 绪论", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    add_body_paragraph(
        doc,
        "高校学生就业指导通常依赖学生自评、教师经验和离散的简历材料。传统流程难以同时回答三个问题：学生当前具备什么能力、"
        "这些能力有哪些可验证证据、哪些岗位既适合当前投递又适合作为下一阶段成长目标。大语言模型能够理解非结构化简历和岗位文本，"
        "但若将全部计算直接交给远程模型，会产生响应时间长、费用不可控、输出结构不稳定和难以复现等工程问题。"
    )
    add_body_paragraph(
        doc,
        "本项目将本地可解释算法与大语言模型协作结合：本地算法承担数据清洗、技能识别、向量召回、基础评分、结果去重和降级兜底；"
        "大模型承担复杂语义理解、四维画像草案和候选岗位精排。系统通过结构化工具调用、证据矩阵和质量审计约束生成边界，"
        "以避免将“智能”简单等同于单次文本生成。"
    )

    doc.add_heading("1.2 问题定义", level=2)
    add_bullet(doc, "能力画像生成链路中多次串行调用 LLM，导致页面等待时间长，且任一调用异常会使整体失败。")
    add_bullet(doc, "岗位匹配入口同步执行大量岗位评分和成长路径生成，用户进入页面时长时间无结果。")
    add_bullet(doc, "岗位库中同义标题和同一岗位族数量较多，TOP5 容易出现 Java 后端岗位重复占位。")
    add_bullet(doc, "仅使用内存缓存时，进程重启后计算结果丢失，重复刷新仍可能触发昂贵计算。")
    add_bullet(doc, "简历 PDF 文本层可能将英文技术词拆分为空格字符，造成姓名、专业、学历和技能抽取不准确。")

    doc.add_heading("1.3 建设目标", level=2)
    add_numbered(doc, "建立可追溯的学生四维能力画像，展示评分、证据、风险提示和智能体协作过程。")
    add_numbered(doc, "将岗位匹配页面首屏改为本地即时结果，AI 精排和成长路径改为按需调用。")
    add_numbered(doc, "建立岗位标题归一化与岗位族多样化机制，提高推荐结果的覆盖面。")
    add_numbered(doc, "通过 MySQL 持久化缓存复用结果，使应用重启后缓存仍然有效。")
    add_numbered(doc, "为简历优化、课程映射、模拟面试和成长趋势提供统一的学生数据基础。")

    doc.add_heading("1.4 主要工作与创新点", level=2)
    add_callout(
        doc,
        "核心思路",
        "使用“本地算法负责速度、稳定和可解释，大模型负责语义增强”的混合架构。系统不是取消工作流，"
        "而是把多个专家的语义生成合并为一次 LLM 请求，再由多个工具节点分别复核和扩展。"
    )
    add_bullet(doc, "单次 LLM、多角色协作：保留五智能体、七次工具调用和五次交接记录。")
    add_bullet(doc, "双向岗位匹配：同时评价学生适岗性与岗位对学生成长的适合度。")
    add_bullet(doc, "多样性约束排序：基于归一化标题、岗位族和分数阈值去重，不强行插入低相关岗位。")
    add_bullet(doc, "版本感知持久化缓存：诊断、岗位库和算法变化时自动形成新缓存键。")
    add_bullet(doc, "按需生成成长路径：用户点击具体岗位后才调用 LLM，避免阻塞岗位列表页面。")


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("2 需求分析", level=1)
    doc.add_heading("2.1 用户与业务场景", level=2)
    add_body_paragraph(
        doc,
        "系统主要面向高校学生和就业指导场景。学生完成注册登录后，可录入基础信息或上传简历，生成能力画像，查看岗位推荐、"
        "技能差距与成长路径，并通过简历优化、课程映射、模拟面试和历史趋势形成连续的就业准备闭环。"
    )

    add_table(
        doc,
        ["角色", "主要诉求", "系统支持"],
        [
            ("学生用户", "快速了解自身能力、岗位适配度和提升方向", "画像、雷达图、TOP5 岗位、成长路径、简历优化、模拟面试"),
            ("就业指导教师", "查看评价依据和学生成长变化", "证据卡、风险提示、历史记录、成长趋势"),
            ("系统维护者", "维护岗位库、模型配置和结果可追溯性", "Excel 导入、环境变量、MySQL 数据表、缓存版本"),
        ],
        [1450, 3100, 4230],
    )
    add_caption(doc, "表 2-1 系统角色与业务诉求")

    doc.add_heading("2.2 功能需求", level=2)
    add_table(
        doc,
        ["编号", "功能模块", "关键功能"],
        [
            ("F01", "用户认证", "注册、登录、退出、基于 Session 的访问控制"),
            ("F02", "学生信息采集", "手工表单录入；简历文本映射到九个数据库字段"),
            ("F03", "能力画像", "四维评分、雷达图、证据卡、风险提示、工作流日志"),
            ("F04", "岗位匹配", "本地召回与评分、TOP5 展示、AI 精排、多样性控制"),
            ("F05", "成长路径", "针对单个岗位按需生成差距、项目建议和三阶段路径"),
            ("F06", "课程映射", "识别课程、课程能力归一、课程到岗位边的可解释评分"),
            ("F07", "简历优化", "PDF/DOCX/TXT 解析、评分、改写建议和行动项"),
            ("F08", "模拟面试", "岗位问题生成、回答评分、反馈和最终总结"),
            ("F09", "历史与成长", "按用户保存多次诊断，展示四维分数变化"),
        ],
        [820, 1900, 6060],
    )
    add_caption(doc, "表 2-2 功能需求清单")

    doc.add_heading("2.3 非功能需求", level=2)
    add_bullet(doc, "性能：进入岗位匹配页面时不等待远程 LLM；缓存命中应达到亚秒级响应。")
    add_bullet(doc, "可靠性：LLM 请求具备超时、重试、结构校验和本地结果兜底机制。")
    add_bullet(doc, "可解释性：所有分数能够关联技能、项目、学历或画像证据，不只输出结论。")
    add_bullet(doc, "一致性：同一诊断、岗位库和算法版本应返回稳定结果；版本变化不得误用旧缓存。")
    add_bullet(doc, "可维护性：业务能力拆分为服务模块，模型参数和数据库连接通过环境变量配置。")
    add_bullet(doc, "隐私与安全：正式部署需采用密码哈希、强会话密钥、HTTPS 和最小化日志原则。")


def add_chapter_3(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("3 系统总体设计", level=1)
    doc.add_heading("3.1 总体架构", level=2)
    add_body_paragraph(
        doc,
        "系统采用分层式单体 Web 架构。表示层由 Jinja2 模板、CSS 和 JavaScript 构成；应用服务层由 FastAPI 路由负责会话、"
        "参数校验、数据库会话和模板渲染；智能与算法层封装能力画像、岗位匹配、课程映射、简历优化和面试服务；"
        "数据层使用 MySQL 保存业务记录和持久化缓存，并通过外部 OpenAI 兼容接口访问大语言模型。"
    )
    add_image(doc, diagrams["architecture"], "图 3-1 系统总体分层架构", 6.15)

    doc.add_heading("3.2 技术选型", level=2)
    add_table(
        doc,
        ["层次", "技术", "选型说明"],
        [
            ("Web 框架", "FastAPI / Starlette", "类型提示友好，支持同步与异步路由，便于构建表单和 JSON 接口[1]"),
            ("模板与前端", "Jinja2 / HTML / CSS / JavaScript", "适合服务端渲染和比赛演示，页面依赖少"),
            ("工作流", "LangGraph", "以有向状态图组织多节点共享状态、工具调用和顺序执行[3]"),
            ("LLM 接入", "LangChain ChatOpenAI", "兼容 DeepSeek 等 OpenAI 协议模型"),
            ("ORM", "SQLAlchemy 2.x", "使用 DeclarativeBase、Mapped 与 Session 管理关系数据[2]"),
            ("数据库", "MySQL / InnoDB", "保存用户、诊断、岗位、日志和持久化缓存[4]"),
            ("检索算法", "本地词袋向量 + 余弦相似度", "无需额外向量数据库，部署成本低"),
            ("测试", "unittest / mock", "覆盖单次 LLM 工作流、快速匹配、岗位多样性和异常路径"),
        ],
        [1450, 2200, 5130],
    )
    add_caption(doc, "表 3-1 核心技术选型")

    doc.add_heading("3.3 模块划分", level=2)
    modules = [
        ("app/main.py", "应用入口、数据模型、路由、数据库会话和模板上下文"),
        ("app/agent/diagnosis_agent.py", "LangGraph 五智能体能力画像工作流"),
        ("app/services/ability_match_service.py", "关键词归一、四维本地评分和岗位规则评分"),
        ("app/services/job_vector_service.py", "岗位稀疏向量索引与余弦召回"),
        ("app/services/llm_ability_match_service.py", "岗位标题归一、多样化、本地匹配与 AI 精排"),
        ("app/services/match_cache_service.py", "岗位版本哈希和缓存算法版本"),
        ("app/services/llm_gap_path_agent.py", "岗位成长路径生成与本地降级"),
        ("app/services/course_job_mapping_service.py", "课程提取、别名归一和课程岗位映射"),
        ("app/services/resume_optimizer_service.py", "PDF/DOCX/TXT 文本提取和简历优化"),
        ("app/services/resume_profile_extractor_service.py", "简历字段结构化抽取"),
    ]
    add_table(doc, ["模块文件", "职责"], modules, [3100, 5680])
    add_caption(doc, "表 3-2 主要代码模块及职责")

    doc.add_heading("3.4 主要业务闭环", level=2)
    add_numbered(doc, "学生登录后录入或上传简历，系统获得结构化学生信息。")
    add_numbered(doc, "LangGraph 工作流形成四维能力画像、证据卡和质量复核结论。")
    add_numbered(doc, "岗位匹配页面读取最新画像，优先使用 MySQL 缓存，否则执行本地快速匹配。")
    add_numbered(doc, "用户按需触发 AI 精排，对本地 TOP10 进行双向语义排序。")
    add_numbered(doc, "用户点击具体岗位后生成成长路径，并写回诊断记录供后续复用。")
    add_numbered(doc, "历史记录与成长趋势按用户聚合多次诊断结果。")


def add_chapter_4(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("4 多智能体能力画像设计", level=1)
    doc.add_heading("4.1 状态图与共享状态", level=2)
    add_body_paragraph(
        doc,
        "能力画像以 DiagnosisState 作为共享状态，保存学生输入、规范化文本、技能、四维分数、证据卡、发展焦点、工具调用日志、"
        "智能体交接记录和质量复核结果。状态图按照画像采集、四维评分、证据抽取、能力归因和质量复核顺序执行，"
        "每个节点仅追加自身负责的结构化产物。"
    )
    add_image(doc, diagrams["workflow"], "图 4-1 单次 LLM 协作与五智能体工作流", 6.25)

    doc.add_heading("4.2 单次 LLM 协作机制", level=2)
    add_body_paragraph(
        doc,
        "原始多专家方案若让四维评分、证据分析、结论生成和质量复核分别调用模型，会形成串行等待。最终代码由“四维评分智能体”"
        "发起一次综合画像请求，一次返回 ability_scores、score_evidence、recognized_skills、assessment_summary、advantages、"
        "weaknesses、dimension_actions 和 quality_notes。后续节点不再重复生成文本，而是调用本地工具复核、扩展和审计共享草案。"
    )
    add_callout(
        doc,
        "协作可见性",
        "一次 LLM 请求不等于单节点流程。系统仍保留五个专家角色、共享工作区、工作流步骤、工具调用记录和交接日志，"
        "前端能够展示“谁调用了什么工具、向谁交接了什么产物”。",
        fill="F0FDF4",
        color=GREEN,
    )

    doc.add_heading("4.3 工具链设计", level=2)
    add_table(
        doc,
        ["工具", "调用角色", "作用"],
        [
            ("ProfileTextNormalizer", "画像采集智能体", "清洗九类学生字段，统计已填与缺失字段"),
            ("SkillKeywordScanner", "画像采集智能体", "识别显性技能并按编程、框架、数据 AI、工程等分组"),
            ("RubricScoreCalculator", "四维评分智能体", "生成可审计的本地评分参考和边界提示"),
            ("LLMAbilityScorer", "四维评分智能体", "一次生成四个专家共用的结构化语义草案"),
            ("EvidenceMatrixTool", "证据抽取智能体", "把分数与证据转为四维证据矩阵和置信度"),
            ("ConsistencyAuditTool", "能力归因智能体", "检查分数、证据与结论之间的冲突"),
            ("FinalAuditTool", "质量复核智能体", "检查工具链、交接记录和岗位排名隔离边界"),
        ],
        [2350, 2100, 4330],
    )
    add_caption(doc, "表 4-1 能力画像工具调用链")

    doc.add_heading("4.4 四维能力模型", level=2)
    add_body_paragraph(
        doc,
        "系统使用专业基础能力、技术实践能力、工具技能能力和职业发展能力四个维度。LLM 评分时需严格引用简历证据；"
        "本地 Rubric 同时计算关键词覆盖、项目/实习/竞赛存在性、工具覆盖和材料完整度，为模型提供边界参考。"
    )
    add_table(
        doc,
        ["维度", "主要评价内容", "典型证据"],
        [
            ("专业基础能力", "专业课程、计算机基础、知识体系、相关证书", "数据结构、算法、数据库、计算机网络、学历与证书"),
            ("技术实践能力", "项目复杂度、实习竞赛、工程落地和量化成果", "项目、部署、性能优化、业务结果、竞赛"),
            ("工具技能能力", "语言、框架、数据库、中间件和部署工具", "Java、Spring Boot、MySQL、Redis、Docker、Git"),
            ("职业发展能力", "目标清晰度、表达、简历完整性和软技能", "目标岗位、自我介绍、沟通协作、材料完整度"),
        ],
        [1700, 3500, 3580],
    )
    add_caption(doc, "表 4-2 四维能力指标说明")

    doc.add_heading("4.5 基础信息与 PDF 文本处理", level=2)
    add_body_paragraph(
        doc,
        "简历字段抽取服务将输入严格映射为 name、major、grade、target_job、skills、projects、competitions、certificates 和 self_intro"
        "九个字段，并拒绝模型添加额外字段。岗位评分中的关键词检测额外处理 PDF 文本层常见的英文字符拆分，例如将"
        "“J a v a”“M y S Q L”“S p r i n g B o o t”压缩后再匹配，从而减少技能漏识别。"
    )


def add_chapter_5(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("5 岗位匹配算法设计", level=1)
    doc.add_heading("5.1 快速路径总体流程", level=2)
    add_body_paragraph(
        doc,
        "岗位匹配采用两阶段架构。第一阶段完全在本地完成召回、规则评分和多样化排序，保证页面能够立即返回；第二阶段由用户按需"
        "触发一次 AI 精排。成长路径与岗位列表分离，只有用户点击某个岗位时才生成对应路径。"
    )
    add_image(doc, diagrams["match_pipeline"], "图 5-1 岗位匹配快速路径与按需 AI 增强", 6.25)

    doc.add_heading("5.2 本地稀疏向量召回", level=2)
    add_body_paragraph(
        doc,
        "系统将学生目标岗位、专业、技能、项目、证书和简历文本拼接为检索文本，并根据目标岗位扩展相关技术词。岗位文本由岗位名、"
        "技能要求、项目、课程、证书、学历、薪资和描述等字段组成。英文按单词切分，中文连续片段生成 2 至 4 字 n-gram，"
        "常见 IT 技术词重复加入以提高权重。词频使用对数缩放："
    )
    add_formula(doc, "w(t) = 1 + ln(tf(t))", "式（5-1）稀疏词袋向量的词频权重")
    add_formula(
        doc,
        "sim(q, d) = (q · d) / (||q||₂ × ||d||₂)",
        "式（5-2）学生向量 q 与岗位向量 d 的余弦相似度",
    )
    add_body_paragraph(
        doc,
        "岗位向量索引按岗位列表签名缓存在进程内，最多保留八个索引版本。缓存命中时复用文本和向量，但重新绑定当前数据库会话"
        "中的 ORM 记录，避免跨请求复用失效对象。"
    )

    doc.add_heading("5.3 本地可解释评分", level=2)
    add_body_paragraph(
        doc,
        "向量召回后的候选岗位使用确定性公式评分。技能覆盖为学生命中的岗位必备技能比例；项目证据来自相关项目关键词；"
        "画像均值由四维能力加权；岗位方向、学历和证书分别提供补充分。"
    )
    add_formula(
        doc,
        "M_local = 55C_skill + 15C_project + 15A_profile + 10F_role + 3F_edu + 2C_cert",
        "式（5-3）本地岗位匹配分；各覆盖率和适配度均归一化到 0～1",
    )
    add_formula(
        doc,
        "A_profile = 0.30P + 0.25R + 0.30T + 0.15C",
        "式（5-4）四维画像加权均值；P、R、T、C 分别代表专业、实践、工具和职业维度",
    )

    doc.add_heading("5.4 岗位去重与多样化", level=2)
    add_body_paragraph(
        doc,
        "系统首先消除大小写、空格、括号说明、职级和常见后缀差异，将“Java后端工程师”“JAVA 后端开发工程师”和带业务说明的"
        "Java 后端标题归一为稳定键。随后根据关键词将岗位归入应用开发、数据、测试、云运维、前端等岗位族。"
    )
    add_bullet(doc, "候选池阶段：目标岗位族最多保留 24 条，其他岗位族每族最多保留 8 条，再按原相似度补齐。")
    add_bullet(doc, "TOPN 阶段：先在分数不低于 60 且与最高分差不超过 18 分的结果中，每族取 1 条；再放宽到每族 2 条。")
    add_bullet(doc, "若多样岗位相关度过低，则按原分数补齐，不为追求多样性强行插入低分岗位。")
    add_callout(
        doc,
        "效果",
        "该策略解决了岗位库中同义岗位大量占位的问题，同时保留“相关性优先”的排序原则。",
        fill="F0FDF4",
        color=GREEN,
    )

    doc.add_heading("5.5 AI 双向精排", level=2)
    add_body_paragraph(
        doc,
        "AI 精排只接收本地 TOP10，并在一次请求中为所有候选输出学生适岗分、岗位适生分、最终匹配分和简短推荐理由。"
        "学生适岗分回答“学生当前能否胜任岗位”，岗位适生分回答“岗位是否适合作为学生下一阶段成长目标”。"
    )
    add_formula(
        doc,
        "M_AI = 0.60S_student→job + 0.40S_job→student",
        "式（5-5）双向 AI 精排总分",
    )
    add_body_paragraph(
        doc,
        "精排客户端默认超时为 120 s，最多重试 3 次，重试间隔采用 1 s、2 s 的指数退避。解析阶段允许单条岗位格式异常，"
        "保留成功解析的部分结果，并用本地候选补齐缺失岗位，避免因一条数据不完整而废弃整批结果。"
    )

    doc.add_heading("5.6 按需成长路径", level=2)
    add_body_paragraph(
        doc,
        "岗位列表页面不再为所有 TOP5 同步生成成长路径。用户点击“查看个性化路径”后，系统只针对一个岗位调用路径智能体，生成"
        "技能差距、推荐项目以及基础补强、项目实践、就业准备三个阶段。结果写入当前诊断记录的 agent_result_json，重复点击时"
        "直接返回已生成路径。若未启用 LLM，也可使用本地规则生成三阶段兜底路径。"
    )


def add_chapter_6(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("6 数据库与持久化缓存设计", level=1)
    doc.add_heading("6.1 数据库总体设计", level=2)
    add_body_paragraph(
        doc,
        "系统使用 SQLAlchemy 声明式模型。启动时通过 Base.metadata.create_all 创建缺失表；正式生产环境建议后续引入 Alembic"
        "管理迁移。当前主应用定义用户、诊断、岗位知识、匹配缓存、数据来源、课程能力、简历和抽取日志等数据结构。"
    )
    add_table(
        doc,
        ["表名", "用途", "当前快照"],
        [
            ("users", "登录账号与创建时间", "1 条"),
            ("diagnosis_records", "学生九类信息、四维分数和智能体 JSON", "18 条"),
            ("job_knowledge_records", "岗位名称、公司、城市、学历、技能、项目、课程、证书和薪资", "679 条"),
            ("job_match_cache_records", "本地排序与 AI 精排持久化缓存", "5 条"),
            ("data_sources / job_post_records", "数据来源与原始岗位采集记录", "当前未填充"),
            ("courses / ability_tags / relation tables", "课程、能力标签及多类关联关系", "当前未填充"),
            ("resume_records / extraction_logs", "简历原文、抽取结果与复核日志", "当前未填充"),
        ],
        [2750, 4440, 1590],
    )
    add_caption(doc, "表 6-1 主要数据表与当前数据库快照")

    doc.add_heading("6.2 诊断记录设计", level=2)
    add_body_paragraph(
        doc,
        "diagnosis_records 以 user_id 关联当前用户，每次提交新增一条记录而不是覆盖旧数据。表中既保存 name、major、grade、"
        "target_job 等结构化字段，也保存四维分数和 agent_result_json。后者承载证据卡、工作流步骤、工具日志、协作记录、"
        "质量复核以及按需生成的成长路径，为历史查看和答辩展示提供完整上下文。"
    )

    doc.add_heading("6.3 MySQL 持久化缓存", level=2)
    add_image(doc, diagrams["cache"], "图 6-1 MySQL 持久化缓存键与失效机制", 6.1)
    add_body_paragraph(
        doc,
        "持久化缓存表 job_match_cache_records 的 cache_key 字段建立唯一索引，diagnosis_id 也建立普通索引。缓存值以 JSON"
        "保存完整 TOP10。系统分别保存 result_type=local 和 result_type=llm 两类结果，并在读取时优先返回 AI 精排缓存。"
    )
    add_formula(
        doc,
        "cache_key = SHA-256(diagnosis_id : job_version : algorithm_version : result_type)",
        "式（6-1）岗位匹配持久化缓存键",
    )
    add_bullet(doc, "diagnosis_id 变化：学生重新诊断后生成新缓存，不误用旧画像。")
    add_bullet(doc, "job_version 变化：岗位数量或技能、课程、证书等字段变化后自动生成新哈希。")
    add_bullet(doc, "algorithm_version 变化：评分公式、Prompt 或召回策略升级时通过版本号主动使旧缓存失效。")
    add_bullet(doc, "result_type 变化：本地排序和 AI 精排互不覆盖，可分别回退和展示。")
    add_body_paragraph(
        doc,
        "缓存保存到 MySQL 的磁盘数据中，因此重启 FastAPI 或 Python 进程只会清空进程内存，不会删除缓存表。应用重新启动并连接"
        "同一数据库后，可以使用相同缓存键直接读取 result_json。若 MySQL 容器没有挂载持久化卷、数据库被清空或版本字段发生变化，"
        "缓存才会失效。"
    )

    doc.add_heading("6.4 数据导入与可追溯性", level=2)
    add_body_paragraph(
        doc,
        "scripts/import_it_jobs_from_xlsx.py 从 data/IT岗位数据.xlsx 导入岗位数据，验证必需列并将技能、项目、课程和证书转换为"
        "JSON 数组。脚本以岗位名称和公司名称判断重复记录，并能为旧表补充公司、城市、学历和薪资字段。"
    )


def add_chapter_7(doc: Document) -> None:
    doc.add_heading("7 其他业务模块设计", level=1)
    doc.add_heading("7.1 简历优化", level=2)
    add_body_paragraph(
        doc,
        "简历优化服务支持 TXT、Markdown、CSV、DOCX 和 PDF。PDF 最多解析前 5 页，并按 pypdf、PyPDF2、pdfplumber 顺序尝试，"
        "以控制处理时间并提高兼容性。优化接口要求模型返回总体分、关键词分、结构分、摘要、优化后简历、优缺点、改写建议和行动项。"
    )

    doc.add_heading("7.2 课程—能力—岗位映射", level=2)
    add_body_paragraph(
        doc,
        "课程映射服务使用课程别名字典，将“数据库系统概论”“数据库原理”等归一为标准课程，再映射到 SQL、MySQL、索引优化、"
        "事务等能力标签。对于本地知识库没有覆盖的课程，可调用 LLM 推理能力标签，并将结果保存到待审核表，而不直接污染真实关系表。"
    )
    add_formula(
        doc,
        "Score_course→job = 50C_ability + 25H_course + 10H_project + 5H_term + 10H_role",
        "式（7-1）课程到岗位的可解释映射评分",
    )

    doc.add_heading("7.3 模拟面试", level=2)
    add_body_paragraph(
        doc,
        "模拟面试服务能够根据目标岗位、岗位描述和简历生成 6 个递进问题，对回答的结构、相关性、证据和表达进行本地评分，"
        "再由 LLM 生成反馈、优点、建议和参考表达。当前 app/main.py 中对应服务导入处于注释状态，因此相关路由虽然已定义，"
        "正式验收前仍需恢复导入并完成端到端回归。"
    )

    doc.add_heading("7.4 历史记录与成长趋势", level=2)
    add_body_paragraph(
        doc,
        "系统按 user_id 查询全部诊断记录，并按时间升序计算专业、实践、工具和职业四维变化。用户既可以查看最近一次画像，也可以"
        "通过记录编号回看历史画像，从而把一次性测评扩展为持续成长档案。"
    )


def add_chapter_8(doc: Document) -> None:
    doc.add_heading("8 性能优化与可靠性设计", level=1)
    doc.add_heading("8.1 性能瓶颈分析", level=2)
    add_body_paragraph(
        doc,
        "优化前的访问记录显示，岗位匹配流程一度约需 415 s。主要原因不是 MySQL 查询本身，而是画像评分、多个岗位批次精排和"
        "TOP5 成长路径生成均在页面请求中串行调用 LLM；此外，缓存检查晚于模型请求，导致重复访问不能提前返回。"
    )

    doc.add_heading("8.2 优化措施", level=2)
    add_table(
        doc,
        ["措施", "实现方式", "预期作用"],
        [
            ("复用四维画像", "直接读取 diagnosis_records 中已保存分数", "避免岗位页再次调用画像评分 LLM"),
            ("本地快速匹配", "稀疏向量召回 TOP30 + 规则评分 TOP10", "进入页面立即显示结果"),
            ("单次 AI 精排", "仅将本地 TOP10 作为一个批次发送", "减少请求数量和输出规模"),
            ("路径按需生成", "点击具体岗位后再调用 LLM", "移除首屏阻塞"),
            ("持久化缓存提前", "所有 LLM 前先查询 MySQL", "重复访问和重启后直接返回"),
            ("延长超时与重试", "120 s 超时、最多 3 次、指数退避", "降低网络抖动导致的误失败"),
            ("结构容错", "单条校验、部分结果保留、本地补齐", "避免整批 JSON 因单条异常失效"),
            ("前端取消 35 s 中断", "等待服务端重试完成", "避免后端仍在执行而浏览器提前放弃"),
        ],
        [2000, 3900, 2880],
    )
    add_caption(doc, "表 8-1 岗位匹配性能与可靠性优化措施")

    doc.add_heading("8.3 实测结果", level=2)
    add_table(
        doc,
        ["场景", "观测结果", "说明"],
        [
            ("优化前岗位匹配访问", "约 415 s", "来自优化过程中的页面访问记录，包含多次串行 LLM"),
            ("真实 TOP10 AI 精排", "约 73.92 s", "DeepSeek 兼容接口；10 条岗位均成功返回 AI 评分"),
            ("AI 精排后刷新页面", "约 178 ms", "直接命中 MySQL llm 缓存"),
            ("岗位数据规模", "679 条", "当前 MySQL 快照"),
            ("核心岗位匹配测试", "10/10 通过", "快速路径、PDF 技术词、向量缓存、精排批次和多样性"),
        ],
        [2700, 2100, 3980],
    )
    add_caption(doc, "表 8-2 当前版本实测数据")
    add_callout(
        doc,
        "解释",
        "73.92 s 是外部模型真实调用的观测值，不是固定服务等级；网络、模型负载和输入长度会影响结果。"
        "系统的关键优化是首屏不等待该调用，并在成功后将结果持久化，使后续访问达到亚秒级。",
        fill="FFF7ED",
        color=AMBER,
    )

    doc.add_heading("8.4 降级与恢复策略", level=2)
    add_bullet(doc, "AI 精排失败时保留本地匹配结果，页面仍可展示岗位。")
    add_bullet(doc, "LLM 返回部分岗位时保留有效结果，并用本地候选补齐。")
    add_bullet(doc, "岗位向量召回无有效相似度时按原岗位顺序保底补齐。")
    add_bullet(doc, "成长路径 LLM 不可用时可生成本地三阶段路径。")
    add_bullet(doc, "MySQL 使用 pool_pre_ping 检测失效连接，避免长时间运行后的连接异常。")


def add_chapter_9(doc: Document) -> None:
    doc.add_heading("9 系统实现与接口", level=1)
    doc.add_heading("9.1 主要接口", level=2)
    add_table(
        doc,
        ["方法", "路径", "功能"],
        [
            ("GET/POST", "/register", "注册页面与账号创建"),
            ("GET/POST", "/login", "登录页面与 Session 建立"),
            ("GET", "/", "登录后的工作台首页"),
            ("POST", "/agent/chat", "上传简历并执行智能体对话分支"),
            ("GET/POST", "/student/input, /student/submit", "学生表单与能力诊断"),
            ("GET", "/ability/profile", "查看最近一次能力画像"),
            ("GET", "/ability/profile/{record_id}", "查看指定历史画像"),
            ("GET", "/job/match", "岗位本地快速匹配与缓存读取"),
            ("POST", "/job/match/refine", "按需 AI 精排 TOP10"),
            ("POST", "/job/match/path", "按需生成单个岗位成长路径"),
            ("GET/POST", "/resume/optimize", "简历解析与优化"),
            ("GET/POST", "/resume/match", "简历课程岗位映射"),
            ("GET/POST", "/interview/mock 等", "模拟面试流程"),
            ("GET", "/history, /growth/trend", "历史诊断和成长趋势"),
            ("GET", "/health", "服务健康检查"),
        ],
        [1300, 3100, 4380],
        font_size=10,
    )
    add_caption(doc, "表 9-1 系统主要 HTTP 接口")

    doc.add_heading("9.2 页面状态设计", level=2)
    add_body_paragraph(
        doc,
        "岗位匹配页面明确区分“本地即时匹配”和“AI 精排结果”。本地结果展示精排按钮；点击后按钮变为“AI 正在精排”并禁用，"
        "完成后页面重新加载，显示“AI 精排已完成”和“MySQL 缓存”状态。该设计避免用户把快速本地结果误认为 AI 输出，"
        "同时对长耗时操作提供可见反馈。"
    )

    doc.add_heading("9.3 配置项", level=2)
    add_table(
        doc,
        ["环境变量", "用途", "当前默认/说明"],
        [
            ("DATABASE_URL", "MySQL SQLAlchemy 连接串", "必填"),
            ("SESSION_SECRET_KEY", "Session 签名密钥", "生产环境必须显式配置强随机值"),
            ("USE_LLM", "是否启用 LLM", "true / false"),
            ("LLM_API_KEY", "OpenAI 兼容接口密钥", "必填且不得提交到版本库"),
            ("LLM_BASE_URL", "模型服务地址", "可配置 DeepSeek 等兼容服务"),
            ("LLM_MODEL", "通用模型名", "各业务可使用专用变量覆盖"),
            ("ABILITY_MATCH_MODEL", "岗位精排模型名", "优先于通用模型"),
            ("JOB_MATCH_LLM_TIMEOUT_SECONDS", "岗位精排单次超时", "默认 120 s"),
            ("JOB_MATCH_LLM_MAX_ATTEMPTS", "岗位精排最大尝试次数", "默认 3"),
            ("JOB_MATCH_MAX_CONCURRENCY", "多批次并发上限", "默认 3"),
            ("INIT_DEMO_JOB_DATA", "是否初始化演示岗位", "默认 false"),
            ("SQLALCHEMY_ECHO", "是否输出 SQL 日志", "默认 false"),
        ],
        [3000, 3500, 2280],
        font_size=10,
    )
    add_caption(doc, "表 9-2 主要环境变量")


def add_chapter_10(doc: Document) -> None:
    doc.add_heading("10 测试与验证", level=1)
    doc.add_heading("10.1 测试方法", level=2)
    add_body_paragraph(
        doc,
        "测试采用 unittest 和 unittest.mock。确定性算法通过固定学生与岗位样本验证；LLM 工作流使用模拟返回验证调用次数和结构；"
        "真实精排则在登录态页面中执行端到端验证，检查按钮状态、结果落库和刷新缓存。"
    )

    doc.add_heading("10.2 核心岗位匹配回归", level=2)
    add_table(
        doc,
        ["测试类别", "验证点", "结果"],
        [
            ("快速路径", "默认岗位匹配不创建 LLM 客户端", "通过"),
            ("成长路径", "默认本地路径包含三个阶段", "通过"),
            ("向量缓存", "缓存向量重新绑定当前请求 ORM 记录", "通过"),
            ("PDF 文本", "空格拆分的 Java/MySQL/Redis/Spring Boot 可识别", "通过"),
            ("AI 精排", "本地 TOP10 以单批次发送", "通过"),
            ("异常降级", "路径 LLM 失败时回退本地路径", "通过"),
            ("标题归一", "Java 后端同义标题映射为统一键", "通过"),
            ("TOP5 多样性", "优先返回不同岗位族且去重", "通过"),
            ("候选池配额", "目标族 24 条、其他族 8 条", "通过"),
            ("相关性边界", "不为多样性强行插入低分岗位", "通过"),
        ],
        [1900, 5000, 1880],
    )
    add_caption(doc, "表 10-1 岗位匹配核心测试结果（10/10 通过）")

    doc.add_heading("10.3 完整测试集现状", level=2)
    add_body_paragraph(
        doc,
        "2026 年 6 月 12 日在当前工作区执行 unittest discover 共发现 16 项测试，其中 12 项通过、4 项报错。"
        "该结果必须与核心岗位匹配的 10/10 通过分开理解。"
    )
    add_table(
        doc,
        ["未通过项", "原因分析", "处理建议"],
        [
            ("简历字段抽取 2 项", "测试直接访问真实 LLM；离线沙箱网络连接失败", "为 _create_llm 注入 FakeLLM，禁止单元测试依赖公网"),
            ("简历优化 2 项", "测试仍要求 response_format 参数和二次重试，但当前服务实现未保持该协议", "同步服务实现与测试契约，明确是否恢复结构化输出重试"),
        ],
        [2150, 3990, 2640],
    )
    add_caption(doc, "表 10-2 完整测试集中尚未闭环的问题")
    add_callout(
        doc,
        "验收结论",
        "岗位匹配核心链路已通过单元测试和真实页面验证；简历字段抽取与简历优化仍需完成测试隔离和协议同步，"
        "因此不应在正式验收报告中宣称“全部测试通过”。",
        fill="FFF1F2",
        color=RED,
    )

    doc.add_heading("10.4 端到端验证", level=2)
    add_numbered(doc, "登录后进入 /job/match，页面立即展示本地 TOP5。")
    add_numbered(doc, "点击“使用 AI 精排 TOP10”，页面进入等待状态且不再在 35 s 主动中断。")
    add_numbered(doc, "约 73.92 s 后返回 10 条 AI 评分，页面展示 AI 精排结果。")
    add_numbered(doc, "结果写入 job_match_cache_records 的 llm 缓存。")
    add_numbered(doc, "重新访问 /job/match，约 178 ms 返回并显示“MySQL 缓存”。")


def add_chapter_11(doc: Document) -> None:
    doc.add_heading("11 部署、运维与安全", level=1)
    doc.add_heading("11.1 运行环境", level=2)
    add_body_paragraph(
        doc,
        "系统运行需要 Python、MySQL 和可选的外部 LLM 服务。项目依赖包括 FastAPI、Uvicorn、SQLAlchemy、PyMySQL、"
        "python-dotenv、LangChain、LangGraph、python-docx 以及 PDF 解析库。requirements.txt 当前包含完整环境导出，"
        "其中混有平台相关依赖，正式交付时建议重新整理最小化依赖文件。"
    )
    add_formula(doc, "python -m uvicorn app.main:app --host 127.0.0.1 --port 8080", "式（11-1）本地启动命令")

    doc.add_heading("11.2 部署步骤", level=2)
    add_numbered(doc, "创建 MySQL 数据库，并在 .env 中配置 DATABASE_URL。")
    add_numbered(doc, "配置 LLM_API_KEY、LLM_BASE_URL、LLM_MODEL 和 USE_LLM。")
    add_numbered(doc, "安装项目依赖，执行岗位 Excel 导入脚本。")
    add_numbered(doc, "启动 Uvicorn；Base.metadata.create_all 自动创建缺失表。")
    add_numbered(doc, "访问 /health 验证服务状态，再注册账号并执行完整业务流程。")

    doc.add_heading("11.3 运维要点", level=2)
    add_bullet(doc, "岗位库或评分算法变更后更新 MATCH_CACHE_ALGORITHM_VERSION，避免复用旧结果。")
    add_bullet(doc, "定期清理不再使用的历史缓存，并监控 job_match_cache_records 的增长。")
    add_bullet(doc, "外部 LLM 应记录耗时、重试次数和失败类型，但日志中不得包含 API 密钥和完整个人简历。")
    add_bullet(doc, "数据库应定期备份，并确保容器部署时挂载 MySQL 数据卷。")

    doc.add_heading("11.4 安全风险", level=2)
    add_callout(
        doc,
        "高优先级风险",
        "当前 User 模型以明文保存密码，且 Session 密钥存在默认值，仅适用于本地演示。生产部署前必须使用 Argon2 或 bcrypt"
        "保存密码哈希，强制配置随机 SESSION_SECRET_KEY，并启用 HTTPS、CSRF 防护、上传文件大小限制和访问审计。",
        fill="FFF1F2",
        color=RED,
    )


def add_chapter_12(doc: Document) -> None:
    doc.add_heading("12 局限性与后续工作", level=1)
    add_body_paragraph(
        doc,
        "当前版本已经完成能力画像、岗位快速匹配、AI 精排和持久化缓存的核心闭环，但仍属于面向本地演示和课程项目的工程版本。"
        "后续工作应优先解决安全、测试隔离和数据治理，再扩展更复杂的模型能力。"
    )
    add_numbered(doc, "认证安全：密码哈希、CSRF、会话过期、登录限流和权限分级。")
    add_numbered(doc, "数据库迁移：引入 Alembic，统一 app/main.py 与 app/models 下的模型定义。")
    add_numbered(doc, "测试工程：所有 LLM 单元测试使用依赖注入和固定响应，端到端测试单独标记。")
    add_numbered(doc, "异步任务：将耗时 AI 精排和路径生成迁移到后台任务或队列，并通过轮询/SSE 展示进度。")
    add_numbered(doc, "向量检索：在岗位规模扩大后升级为 embedding 与专用向量索引，同时保留可解释规则分。")
    add_numbered(doc, "扫描件简历：接入 OCR，并对姓名、专业、学历等字段增加规则与人工确认界面。")
    add_numbered(doc, "推荐评估：建设人工标注集，使用 NDCG、Recall@K、多样性和覆盖率等指标评估排序。")
    add_numbered(doc, "数据治理：填充 data_sources、extraction_logs 和关系表，形成来源、模型版本和人工复核链路。")
    add_numbered(doc, "模拟面试：恢复主程序服务导入，补充会话持久化、异常降级和完整回归测试。")


def add_conclusion(doc: Document) -> None:
    doc.add_heading("结论", level=1)
    add_body_paragraph(
        doc,
        "本文基于当前最终工作区代码，对岗位能力达成学生成长诊断与精准就业智能体系统进行了完整技术说明。系统通过 LangGraph"
        "组织五智能体协作流程，以一次 LLM 请求生成共享语义草案，再由本地工具完成证据、归因和质量复核；岗位推荐采用稀疏向量召回、"
        "可解释评分、岗位族多样化和按需 AI 双向精排；MySQL 持久化缓存使本地和 AI 结果能够跨进程重启复用。"
    )
    add_body_paragraph(
        doc,
        "实测表明，真实 TOP10 AI 精排能够稳定完成，结果持久化后页面刷新进入亚秒级。该架构在不牺牲工作流展示和语义能力的前提下，"
        "将高成本模型调用从页面首屏剥离，形成了速度、可靠性、可解释性和智能增强之间较为平衡的实现。与此同时，完整测试隔离、密码安全、"
        "数据库迁移和后台任务仍是正式生产部署前必须完成的工作。"
    )


def add_references(doc: Document) -> None:
    doc.add_heading("参考文献", level=1)
    references = [
        ("[1] FastAPI Project. FastAPI Documentation [EB/OL]. ", "https://fastapi.tiangolo.com/"),
        ("[2] SQLAlchemy Authors. SQLAlchemy 2.0 Documentation [EB/OL]. ", "https://docs.sqlalchemy.org/en/20/"),
        ("[3] LangChain Inc. LangGraph Overview [EB/OL]. ", "https://docs.langchain.com/oss/python/langgraph/overview"),
        ("[4] Oracle Corporation. MySQL 8.4 Reference Manual [EB/OL]. ", "https://dev.mysql.com/doc/refman/8.4/en/"),
        (
            "[5] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need [C]//"
            "Advances in Neural Information Processing Systems. 2017: 5998-6008. ",
            "https://arxiv.org/abs/1706.03762",
        ),
        ("[6] Python Software Foundation. asyncio — Asynchronous I/O [EB/OL]. ", "https://docs.python.org/3/library/asyncio.html"),
        ("[7] Pallets. Jinja Documentation [EB/OL]. ", "https://jinja.palletsprojects.com/"),
        ("[8] National Institute of Standards and Technology. Secure Hash Standard (FIPS PUB 180-4) [S]. ", "https://doi.org/10.6028/NIST.FIPS.180-4"),
    ]
    for prefix, url in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(-24)
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.35
        set_run_font(paragraph.add_run(prefix), size=10.5)
        add_hyperlink(paragraph, url, url)


def add_appendices(doc: Document) -> None:
    doc.add_heading("附录 A 关键源代码索引", level=1)
    add_table(
        doc,
        ["主题", "文件路径", "关键对象"],
        [
            ("应用入口", "app/main.py", "FastAPI、ORM 模型、路由、持久化缓存读写"),
            ("画像工作流", "app/agent/diagnosis_agent.py", "DiagnosisState、build_diagnosis_graph、run_diagnosis_agent"),
            ("本地画像评分", "app/services/ability_match_service.py", "score_four_dimensions、match_profile_to_job"),
            ("向量召回", "app/services/job_vector_service.py", "build_text_vector、cosine_similarity、retrieve_jobs_by_vector"),
            ("岗位匹配", "app/services/llm_ability_match_service.py", "diversify_job_matches、calculate_local_job_match、refine_job_matches_with_llm"),
            ("缓存版本", "app/services/match_cache_service.py", "build_job_version、MATCH_CACHE_ALGORITHM_VERSION"),
            ("成长路径", "app/services/llm_gap_path_agent.py", "generate_top5_gap_paths"),
            ("课程映射", "app/services/course_job_mapping_service.py", "extract_courses_from_resume、build_course_job_mapping_graph"),
            ("岗位导入", "scripts/import_it_jobs_from_xlsx.py", "import_jobs、ensure_extra_columns"),
        ],
        [1800, 3700, 3280],
        font_size=9.5,
    )
    add_caption(doc, "表 A-1 关键源代码位置")

    doc.add_heading("附录 B 最终版本验收检查表", level=1)
    checks = [
        ("能力画像工作流", "完成", "单次 LLM、五角色、七工具、五次交接"),
        ("基础信息持久化", "完成", "姓名、专业、年级、目标岗位写入 diagnosis_records"),
        ("能力雷达图", "完成", "前端脚本从四维分数渲染"),
        ("岗位本地快速匹配", "完成", "进入页面不创建 LLM 客户端"),
        ("推荐多样性", "完成", "标题归一、岗位族配额和相关性阈值"),
        ("AI TOP10 精排", "完成", "120 s 超时、3 次尝试、部分容错"),
        ("MySQL 持久化缓存", "完成", "应用重启后可复用"),
        ("单岗位成长路径", "完成", "点击后按需生成并写回诊断 JSON"),
        ("岗位匹配核心测试", "完成", "10/10 通过"),
        ("完整测试集", "待完善", "12/16 通过，4 项需隔离外部依赖或同步契约"),
        ("生产安全", "待完善", "密码哈希、强 Session 密钥、迁移和 HTTPS"),
        ("模拟面试端到端", "待完善", "服务实现存在，主程序导入需恢复"),
    ]
    add_table(doc, ["检查项", "状态", "说明"], checks, [2700, 1450, 4630], font_size=10)
    add_caption(doc, "表 B-1 最终版本验收状态")

    doc.add_heading("附录 C 文档基准说明", level=1)
    add_body_paragraph(
        doc,
        "本技术文档以 2026 年 6 月 12 日工作区代码为基线，并以 MySQL 快照、自动化测试和端到端页面验证为依据。"
        "规划项与尚未闭环的内容均已明确标注，不将设想表述为已完成成果。"
    )


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()
    doc = Document()
    configure_styles(doc)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_cover(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, cover=False)
    restart_page_number(body_section, 1)

    add_abstracts(doc)
    add_toc(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc, diagrams)
    add_chapter_4(doc, diagrams)
    add_chapter_5(doc, diagrams)
    add_chapter_6(doc, diagrams)
    add_chapter_7(doc)
    add_chapter_8(doc)
    add_chapter_9(doc)
    add_chapter_10(doc)
    add_chapter_11(doc)
    add_chapter_12(doc)
    add_conclusion(doc)
    add_references(doc)
    add_appendices(doc)

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal" and paragraph.text.strip():
            paragraph.paragraph_format.widow_control = True

    doc.core_properties.title = "岗位能力达成学生成长诊断与精准就业智能体系统技术文档"
    doc.core_properties.subject = "系统设计、实现、算法、数据库、性能优化与测试"
    doc.core_properties.author = "job-ability-agent-system 项目组"
    doc.core_properties.keywords = "学生能力画像, 多智能体, 岗位匹配, MySQL缓存, LLM"
    doc.core_properties.comments = "基于2026-06-12最终工作区代码生成"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    result = build_document()
    print(result)
