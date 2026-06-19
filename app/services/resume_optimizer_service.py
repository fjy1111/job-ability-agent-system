from __future__ import annotations

import json
import os
import re
import hashlib
from collections import OrderedDict
from copy import deepcopy
from html import escape as html_escape
from io import BytesIO
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from langchain_openai import ChatOpenAI

from app.services.resume_expert_kb_service import (
    load_resume_expert_rules,
    retrieve_resume_expert_rules,
)
from app.services.llm_errors import LLMCallError

load_dotenv()


TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
MAX_PDF_PAGES = 5
RESUME_OPTIMIZER_CACHE_VERSION = "resume_optimizer_v2"
RESUME_OPTIMIZER_CACHE_MAX_SIZE = 64
RESUME_OPTIMIZER_EXPERT_RULE_LIMIT = 12

_resume_optimizer_cache: OrderedDict[str, dict[str, Any]] = OrderedDict()


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _first_pdf_pages(pages):
    for index, page in enumerate(pages):
        if index >= MAX_PDF_PAGES:
            break
        yield page


def _extract_docx_text(file_content: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ""

    document = Document(BytesIO(file_content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)
    return "\n".join([*paragraphs, *table_cells]).strip()


def _extract_pdf_text(file_content: bytes) -> tuple[str, str]:
    """
    多解析器兜底提取 PDF 文本。
    优先 pypdf，其次 PyPDF2，再其次 pdfplumber；只读取前几页以避免上传后长时间等待。
    """
    missing_parsers: list[str] = []

    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_content))
        text = "\n".join(page.extract_text() or "" for page in _first_pdf_pages(reader.pages))
        if text.strip():
            return text.strip(), ""
    except ModuleNotFoundError:
        missing_parsers.append("pypdf")
    except Exception as exc:
        return "", f"PDF 解析失败：{type(exc).__name__}，请粘贴简历文本。"

    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(file_content))
        text = "\n".join(page.extract_text() or "" for page in _first_pdf_pages(reader.pages))
        if text.strip():
            return text.strip(), ""
    except ModuleNotFoundError:
        missing_parsers.append("PyPDF2")
    except Exception:
        pass

    try:
        import pdfplumber

        with pdfplumber.open(BytesIO(file_content)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in _first_pdf_pages(pdf.pages))
        if text.strip():
            return text.strip(), ""
    except ModuleNotFoundError:
        missing_parsers.append("pdfplumber")
    except Exception:
        pass

    if missing_parsers:
        return "", "PDF 解析依赖未安装，请在 agent 环境执行：python -m pip install pypdf，然后重启项目。"

    return "", "PDF 未提取到有效文本，可能是扫描件或图片版简历，请粘贴简历文本。"


def extract_resume_text_from_upload(filename: str, file_content: bytes) -> tuple[str, list[str]]:
    """
    从上传文件中提取简历文本。只负责文件解析，不生成替代优化结果。
    """
    suffix = Path(filename).suffix.lower()
    warnings: list[str] = []

    if suffix in TEXT_EXTENSIONS:
        return _decode_text(file_content).strip(), warnings

    if suffix in DOCX_EXTENSIONS:
        text = _extract_docx_text(file_content)
        if text:
            return text, warnings
        warnings.append("Word 文件暂时无法解析，请粘贴简历文本。")
        return "", warnings

    if suffix in PDF_EXTENSIONS:
        text, parser_warning = _extract_pdf_text(file_content)
        if text:
            return text, warnings
        warnings.append(parser_warning)
        return "", warnings

    warnings.append("该文件暂时无法解析，请粘贴简历文本。")
    return "", warnings


def _safe_json_loads(text: str) -> dict[str, Any]:
    if not text:
        raise LLMCallError()

    text = text.strip()
    text = re.sub(r"^```(?:json)?", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(r"```$", "", text).strip()

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", text)
        if not match:
            raise LLMCallError()
        try:
            data = json.loads(match.group(0))
        except json.JSONDecodeError:
            raise LLMCallError()

    if not isinstance(data, dict):
        raise LLMCallError()
    return data


def _create_llm() -> ChatOpenAI:
    if os.getenv("USE_LLM", "true").lower() != "true":
        raise LLMCallError()

    api_key = (
        os.getenv("LLM_API_KEY")
        or os.getenv("OPENAI_API_KEY")
        or os.getenv("DEEPSEEK_API_KEY")
        or os.getenv("DASHSCOPE_API_KEY")
    )
    model = (
        os.getenv("RESUME_OPTIMIZER_MODEL")
        or os.getenv("LLM_MODEL")
        or os.getenv("OPENAI_MODEL")
        or os.getenv("DEEPSEEK_MODEL")
        or os.getenv("DASHSCOPE_MODEL")
    )
    base_url = (
        os.getenv("LLM_BASE_URL")
        or os.getenv("OPENAI_BASE_URL")
        or os.getenv("DEEPSEEK_BASE_URL")
        or os.getenv("DASHSCOPE_BASE_URL")
    )

    if not api_key or not model:
        raise LLMCallError()

    kwargs: dict[str, Any] = {
        "model": model,
        "api_key": api_key,
        "temperature": 0.25,
        "timeout": 60,
        "max_retries": 0,
    }
    if base_url:
        kwargs["base_url"] = base_url
    return ChatOpenAI(**kwargs)


def _score(value: Any) -> int:
    try:
        number = int(round(float(value)))
    except (TypeError, ValueError):
        raise LLMCallError()
    return max(0, min(100, number))


def _list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _validate_result(data: dict[str, Any]) -> dict[str, Any]:
    required = [
        "overall_score",
        "keyword_score",
        "structure_score",
        "summary",
        "optimized_resume",
        "strengths",
        "weaknesses",
        "rewrite_suggestions",
        "action_items",
    ]
    if any(key not in data for key in required):
        raise LLMCallError()

    rewrite_suggestions = _list(data.get("rewrite_suggestions"))
    if not all(isinstance(item, dict) for item in rewrite_suggestions):
        raise LLMCallError()

    return {
        "overall_score": _score(data.get("overall_score")),
        "keyword_score": _score(data.get("keyword_score")),
        "structure_score": _score(data.get("structure_score")),
        "summary": str(data.get("summary", "")).strip(),
        "optimized_resume": str(data.get("optimized_resume", "")).strip(),
        "strengths": [str(item).strip() for item in _list(data.get("strengths")) if str(item).strip()],
        "weaknesses": [str(item).strip() for item in _list(data.get("weaknesses")) if str(item).strip()],
        "rewrite_suggestions": rewrite_suggestions,
        "action_items": [str(item).strip() for item in _list(data.get("action_items")) if str(item).strip()],
        "agent_warning": "",
    }


def _public_expert_rule(rule: dict[str, Any]) -> dict[str, str]:
    return {
        "id": str(rule.get("id", "")).strip(),
        "title": str(rule.get("title", "")).strip(),
        "category": str(rule.get("category", "")).strip(),
        "suggestion": str(rule.get("suggestion", "")).strip(),
        "source_image": str(rule.get("source_image", "")).strip(),
    }


def _expert_rules_prompt(matched_rules: list[dict]) -> str:
    if not matched_rules:
        return ""

    rule_lines = []
    for index, rule in enumerate(matched_rules, start=1):
        title = str(rule.get("title", "")).strip()
        suggestion = str(rule.get("suggestion", "")).strip()
        if title and suggestion:
            rule_lines.append(f"{index}. {title}：{suggestion}")
        elif suggestion:
            rule_lines.append(f"{index}. {suggestion}")

    if not rule_lines:
        return ""

    return f"""
【专家手工标注建议知识库】
以下是从历史专家批注中检索到的规则，请优先参考：

{chr(10).join(rule_lines)}

生成简历优化建议时，请结合学生简历内容和上述专家规则，重点检查：
* 是否缺少岗位名称
* 是否缺少项目时间
* 是否缺少个人角色
* 项目功能是否描述清楚
* 技术亮点是否写清楚“技术方案 + 解决问题 + 效果”
* 是否有量化成果
* 技能描述是否过泛
* 教育经历和证书表达是否规范
"""


def _rule_identity(rule: dict[str, Any]) -> str:
    return str(rule.get("id") or rule.get("title") or "").strip()


def _retrieve_optimizer_expert_rules(
    resume_text: str,
    max_rules: int = RESUME_OPTIMIZER_EXPERT_RULE_LIMIT,
) -> list[dict]:
    selected = retrieve_resume_expert_rules(resume_text, max_rules=max_rules)
    seen = {_rule_identity(rule) for rule in selected if _rule_identity(rule)}

    if len(selected) >= max_rules:
        return selected[:max_rules]

    supplemental = sorted(
        load_resume_expert_rules(),
        key=lambda item: item.get("priority", 0),
        reverse=True,
    )
    for rule in supplemental:
        identity = _rule_identity(rule)
        if not identity or identity in seen:
            continue
        selected.append(rule)
        seen.add(identity)
        if len(selected) >= max_rules:
            break
    return selected[:max_rules]


def _resume_optimizer_cache_key(
    resume_text: str,
    job_description: str,
    target_role: str,
    output_language: str,
    harvard_format: bool,
    expert_rules: list[dict],
) -> str:
    payload = {
        "version": RESUME_OPTIMIZER_CACHE_VERSION,
        "resume_text": str(resume_text or "").strip(),
        "job_description": str(job_description or "").strip(),
        "target_role": str(target_role or "").strip(),
        "output_language": str(output_language or "").strip(),
        "harvard_format": bool(harvard_format),
        "expert_rule_ids": [_rule_identity(rule) for rule in expert_rules],
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def clear_resume_optimizer_cache() -> None:
    _resume_optimizer_cache.clear()


def _get_cached_resume_result(cache_key: str) -> dict[str, Any] | None:
    cached = _resume_optimizer_cache.get(cache_key)
    if cached is None:
        return None
    _resume_optimizer_cache.move_to_end(cache_key)
    result = deepcopy(cached)
    result["cache_hit"] = True
    return result


def _save_resume_result_to_cache(cache_key: str, result: dict[str, Any]) -> None:
    cached = deepcopy(result)
    cached["cache_hit"] = False
    _resume_optimizer_cache[cache_key] = cached
    _resume_optimizer_cache.move_to_end(cache_key)
    while len(_resume_optimizer_cache) > RESUME_OPTIMIZER_CACHE_MAX_SIZE:
        _resume_optimizer_cache.popitem(last=False)


def _rule_text(rule: dict[str, Any]) -> str:
    patterns = rule.get("problem_patterns")
    if isinstance(patterns, list):
        pattern_text = " ".join(str(item) for item in patterns)
    else:
        pattern_text = str(patterns or "")
    return " ".join(
        str(rule.get(key, "") or "")
        for key in ("title", "category", "suggestion")
    ) + f" {pattern_text}"


def _find_rule_title(
    expert_rules_used: list[dict[str, Any]],
    keywords: tuple[str, ...],
) -> str:
    for rule in expert_rules_used:
        haystack = _rule_text(rule).casefold()
        if any(keyword.casefold() in haystack for keyword in keywords):
            title = str(rule.get("title", "")).strip()
            if title:
                return title
    return ""


def _add_highlight_matches(
    matches: list[tuple[int, int, str]],
    text: str,
    patterns: list[str],
    title: str,
) -> None:
    if not title:
        return

    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            start, end = match.span()
            while start < end and text[start].isspace():
                start += 1
            while end > start and text[end - 1].isspace():
                end -= 1
            if start < end:
                matches.append((start, end, title))


def _select_non_overlapping_matches(
    matches: list[tuple[int, int, str]],
) -> list[tuple[int, int, str]]:
    selected: list[tuple[int, int, str]] = []
    cursor = 0
    for start, end, title in sorted(matches, key=lambda item: (item[0], -(item[1] - item[0]))):
        if start < cursor:
            continue
        selected.append((start, end, title))
        cursor = end
    return selected


def _add_new_text_highlight_matches(
    matches: list[tuple[int, int, str]],
    original_text: str,
    optimized_text: str,
    patterns: list[str],
    title: str,
) -> None:
    if not title:
        return

    original_compact = re.sub(r"\s+", "", str(original_text or "")).casefold()
    for pattern in patterns:
        for match in re.finditer(pattern, optimized_text, flags=re.IGNORECASE):
            start, end = match.span()
            while start < end and optimized_text[start].isspace():
                start += 1
            while end > start and optimized_text[end - 1].isspace():
                end -= 1
            if start >= end:
                continue

            phrase = optimized_text[start:end]
            phrase_compact = re.sub(r"\s+", "", phrase).casefold()
            if phrase_compact and phrase_compact not in original_compact:
                matches.append((start, end, title))


def build_highlighted_optimized_html(
    original_resume_text: str,
    optimized_resume_text: str,
    expert_rules_used: list[dict],
) -> str:
    """
    Build safe HTML for the optimized resume and highlight text likely added
    because of expert annotation rules.
    """
    original_text = str(original_resume_text or "")
    text = str(optimized_resume_text or "")
    rules = [rule for rule in (expert_rules_used or []) if isinstance(rule, dict)]
    if not text:
        return ""
    if not rules:
        return html_escape(text, quote=False)

    matches: list[tuple[int, int, str]] = []

    quant_title = _find_rule_title(
        rules,
        ("量化", "提升", "降低", "减少", "响应时间", "命中率", "效果"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:响应时间|接口响应(?:时间|速度)?|人工处理时间|处理时间|缓存命中率|命中率|重复提交率|提交率|查询效率|检索效率|点击率|准确率|性能|效率|耗时|延迟)[^，。；,\n]{0,24}(?:降低|下降|减少|提升|提高|达|达到|降至|低至|缩短|优化至)[^，。；,\n]{0,12}\d+(?:\.\d+)?%",
            r"(?:降低|下降|减少|提升|提高|达|达到|降至|低至|缩短|优化至)[^，。；,\n]{0,12}\d+(?:\.\d+)?%",
            r"(?:降低|下降|减少|提升|提高|优化)[^，。；,\n]{2,30}(?:访问|速度|效率|性能|开销|耗时|响应|一致性)",
        ],
        quant_title,
    )

    ranking_title = _find_rule_title(rules, ("排名", "Top", "GPA")) or quant_title
    _add_highlight_matches(
        matches,
        text,
        [r"Top\s*\d+(?:\.\d+)?%"],
        ranking_title,
    )

    role_title = _find_rule_title(
        rules,
        ("个人角色", "个人贡献", "个人职责", "主要负责", "负责人", "核心成员"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:后端|前端|算法|测试|数据|全栈)?(?:核心成员|核心开发|负责人|开发负责人|开发成员)",
            r"(?:核心后端开发|后端核心开发|后端负责人|个人贡献)",
            r"(?:独立完成|主要负责|本人负责|个人负责)",
            r"作为[^，。；,\n]{0,12}(?:负责人|核心成员|开发成员)",
        ],
        role_title,
    )

    redis_title = _find_rule_title(
        rules,
        ("Redis", "Redisson", "缓存", "分布式锁", "高频查询"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:使用|通过|结合|基于|引入|利用)?\s*(?:Redis|Redisson)[^，。；,\n]{0,40}(?:缓存|分布式锁|一致性|高频查询|读多写少|查询结果|热点数据|缓存命中)[^，。；,\n]{0,30}",
        ],
        redis_title,
    )

    microservice_title = _find_rule_title(
        rules,
        ("微服务", "Spring Cloud", "Spring Cloud Alibaba", "OpenFeign", "服务拆分"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:基于|使用|通过|采用)?\s*Spring Cloud Alibaba[^，。；,\n]{0,30}(?:微服务|生态|架构)[^，。；,\n]{0,20}",
            r"(?:微服务)(?:架构|生态|项目|体系)?",
            r"(?:拆分|划分)[^，。；,\n]{0,30}(?:服务|模块)",
            r"OpenFeign[^，。；,\n]{0,30}(?:服务调用|远程调用|服务间调用)",
        ],
        microservice_title,
    )

    gateway_title = _find_rule_title(
        rules,
        ("Nacos", "Gateway", "注册配置中心", "统一网关", "路由", "跨域"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"Nacos[^，。；,\n]{0,30}(?:注册配置中心|服务注册|配置管理)[^，。；,\n]{0,20}",
            r"Gateway[^，。；,\n]{0,40}(?:统一网关|路由转发|跨域处理|权限入口)[^，。；,\n]{0,20}",
            r"(?:注册配置中心|统一网关|路由转发|跨域处理)",
        ],
        gateway_title,
    )

    security_title = _find_rule_title(
        rules,
        ("Spring Security", "JWT", "认证", "授权", "接口安全", "权限控制"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:Spring Security\s*\+\s*JWT|Spring Security|JWT)[^，。；,\n]{0,40}(?:认证授权|登录认证|接口鉴权|用户身份校验|权限控制|接口安全)[^，。；,\n]{0,20}",
            r"(?:认证授权|登录认证|接口鉴权|用户身份校验|权限控制|接口安全)",
        ],
        security_title,
    )

    xxl_job_title = _find_rule_title(
        rules,
        ("XXL-JOB", "定时任务", "分片广播", "订单状态", "自动处理"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"XXL-JOB[^，。；,\n]{0,50}(?:分片广播|定时任务|订单状态|自动处理|状态同步|超时订单|异常订单)[^，。；,\n]{0,30}",
            r"(?:分片广播|订单状态|自动处理|状态同步|超时订单关闭|异常订单处理)",
        ],
        xxl_job_title,
    )

    docker_title = _find_rule_title(
        rules,
        ("Docker Compose", "Docker", "部署", "运维", "测试环境"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"Docker Compose[^，。；,\n]{0,50}(?:简化|部署|测试环境|环境搭建|复现|部署效率)[^，。；,\n]{0,30}",
            r"(?:简化测试环境搭建|提高部署效率|方便测试复现)",
        ],
        docker_title,
    )

    business_title = _find_rule_title(
        rules,
        ("业务场景", "订单", "配送", "商家", "用户", "调度"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:面向|覆盖|支撑)[^，。；,\n]{0,40}(?:订单|配送|商家|用户|调度)[^，。；,\n]{0,30}(?:场景|业务|问题|流程)",
            r"(?:订单状态流转|配送调度|商家接单|用户下单|履约跟踪)",
        ],
        business_title,
    )

    english_title = _find_rule_title(
        rules,
        ("英语", "四六级", "CET", "分数"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:CET[-\s]?[46]|英语[四六46]级|大学英语[四六46]级)[：:\s-]*\d{3,4}\s*分?",
            r"\d{3,4}\s*分",
        ],
        english_title,
    )

    project_time_title = _find_rule_title(
        rules,
        ("项目时间", "起止时间", "开发周期"),
    )
    _add_new_text_highlight_matches(
        matches,
        original_text,
        text,
        [
            r"(?:项目时间|开发周期)[：:\s]*\d{4}[./-]\d{1,2}\s*(?:-|–|—|~|至)\s*\d{4}[./-]\d{1,2}",
            r"\d{4}[./-]\d{1,2}\s*(?:-|–|—|~|至)\s*\d{4}[./-]\d{1,2}",
        ],
        project_time_title,
    )

    project_detail_title = _find_rule_title(
        rules,
        ("项目功能", "技术亮点", "技术方案", "解决问题", "个人职责与成果"),
    )
    _add_highlight_matches(
        matches,
        text,
        [
            r"(?:个人职责与成果|技术方案\s*\+\s*解决问题\s*\+\s*效果)",
        ],
        project_detail_title,
    )

    selected_matches = _select_non_overlapping_matches(matches)
    if not selected_matches:
        return html_escape(text, quote=False)

    html_parts: list[str] = []
    cursor = 0
    for start, end, title in selected_matches:
        html_parts.append(html_escape(text[cursor:start], quote=False))
        safe_text = html_escape(text[start:end], quote=False)
        safe_title = html_escape(title, quote=True)
        html_parts.append(
            f'<span class="expert-highlight" title="{safe_title}">{safe_text}</span>'
        )
        cursor = end
    html_parts.append(html_escape(text[cursor:], quote=False))
    return "".join(html_parts)


def _attach_expert_rules(
    result: dict[str, Any],
    matched_rules: list[dict],
    original_resume_text: str = "",
) -> dict[str, Any]:
    public_rules = [
        rule
        for rule in (_public_expert_rule(item) for item in matched_rules)
        if rule["title"] or rule["suggestion"]
    ]
    result["expert_rules_used_count"] = len(public_rules)
    result["expert_rules_used"] = public_rules
    try:
        result["highlighted_optimized_html"] = build_highlighted_optimized_html(
            original_resume_text,
            str(result.get("optimized_resume", "")),
            matched_rules,
        )
    except Exception:
        result["highlighted_optimized_html"] = ""
    return result


def _invoke_optimizer_with_retry(prompt: str) -> dict[str, Any]:
    llm = _create_llm()
    prompts = [
        prompt,
        (
            f"{prompt}\n\n"
            "上一次的输出未通过 JSON 结构校验。请严格按要求返回单个 JSON 对象，"
            "不要输出 Markdown、解释文字或额外字段。"
        ),
    ]

    for index, current_prompt in enumerate(prompts):
        try:
            response = llm.invoke(
                current_prompt,
                response_format={"type": "json_object"},
            )
            return _validate_result(_safe_json_loads(str(response.content)))
        except LLMCallError:
            if index == len(prompts) - 1:
                raise

    raise LLMCallError()


def optimize_resume(
    resume_text: str,
    job_description: str,
    target_role: str = "",
    output_language: str = "auto",
    harvard_format: bool = False,
    use_cache: bool = True
) -> dict[str, Any]:
    """
    调用大模型生成简历优化结果。失败时只抛出“调用LLM失败”。
    """
    role = target_role.strip() or "目标岗位"
    matched_rules = _retrieve_optimizer_expert_rules(
        resume_text,
        max_rules=RESUME_OPTIMIZER_EXPERT_RULE_LIMIT,
    )
    cache_key = _resume_optimizer_cache_key(
        resume_text,
        job_description,
        target_role,
        output_language,
        harvard_format,
        matched_rules,
    )
    if use_cache:
        cached_result = _get_cached_resume_result(cache_key)
        if cached_result is not None:
            return cached_result

    expert_rules_text = _expert_rules_prompt(matched_rules)
    prompt = f"""
你是资深就业辅导与简历优化智能体。请基于候选人简历和目标岗位生成可直接展示的简历优化结果。

要求：
1. 不编造候选人没有提供的经历。
2. 输出语言：{output_language}。如果为 auto，请跟随简历主要语言。
3. Harvard 格式要求：{"需要" if harvard_format else "不需要"}。
4. 分数必须是 0-100 整数。
5. 返回严格 JSON，不要 Markdown。

目标岗位：{role}
岗位信息：
{job_description[:3000] or "未提供岗位描述，请做通用求职简历优化。"}

候选人简历：
{resume_text[:6000]}

{expert_rules_text}

JSON 格式：
{{
  "overall_score": 0,
  "keyword_score": 0,
  "structure_score": 0,
  "summary": "100字以内总结",
  "optimized_resume": "完整优化稿或重点段落优化稿",
  "strengths": ["优势1", "优势2", "优势3"],
  "weaknesses": ["待提升1", "待提升2", "待提升3"],
  "rewrite_suggestions": [
    {{
      "section": "项目经历",
      "before": "原写法摘要",
      "after": "建议写法",
      "reason": "修改原因"
    }}
  ],
  "action_items": ["下一步动作1", "下一步动作2", "下一步动作3"]
}}
"""

    try:
        result = _invoke_optimizer_with_retry(prompt)
        result = _attach_expert_rules(result, matched_rules, resume_text)
        result["cache_hit"] = False
        if use_cache:
            _save_resume_result_to_cache(cache_key, result)
        return result
    except LLMCallError:
        raise
    except Exception:
        raise LLMCallError()
